# -*- coding: UTF-8 -*-
from django.shortcuts import render
from django.shortcuts import render_to_response, redirect
from django.template import RequestContext
from django.views.generic import ListView, DetailView, CreateView, UpdateView
from teacher.models import *
from student.models import *
from account.models import Domain, Level, Parent, Log, Message, MessagePoll, MessageContent
from .forms import *
from django.contrib.auth.models import Group
from django.core.exceptions import ObjectDoesNotExist,  MultipleObjectsReturned
from django.core.files.storage import FileSystemStorage
from django.http import HttpResponseRedirect
import os
from uuid import uuid4
from django.conf import settings
from wsgiref.util import FileWrapper
from django.http import HttpResponse
from django.http import JsonResponse
import re
from django.contrib.auth.models import User
from django.db.models import Q
import ast
from docx import *
from docx.shared import Inches
from docx.shared import RGBColor
from django.utils import timezone
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import StringIO
from shutil import copyfile
import xlsxwriter
from django.utils.timezone import localtime
from datetime import datetime
import docx 
import os.path
from django.utils.dateparse import parse_date
from random import shuffle
from operator import attrgetter
from helper import VideoLogHelper
# 判斷是否為授課教師
def is_teacher(user, classroom_id):
    return user.groups.filter(name='teacher').exists() and Classroom.objects.filter(teacher_id=user.id, id=classroom_id).exists()

def is_assistant(user, classroom_id):
    assistants = Assistant.objects.filter(classroom_id=classroom_id, user_id=user.id)
    if len(assistants)>0 :
        return True
    return False	
	
def is_event_open(request):
		return True

# 列出所有課程
class ClassroomListView(ListView):
    model = Classroom
    context_object_name = 'classrooms'
    template_name = 'teacher/classroom.html'
    paginate_by = 25
    def get_queryset(self):      
        queryset = Classroom.objects.filter(teacher_id=self.request.user.id).order_by("-id")
        return queryset
        
#新增一個課程
class ClassroomCreateView(CreateView):
    model = Classroom
    form_class = ClassroomForm
    template_name = 'teacher/classroom_form.html'    
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.teacher_id = self.request.user.id
        self.object.domains = self.request.POST.getlist('domains')
        self.object.levels = self.request.POST.getlist('levels')	
        self.object.save()
        # 將教師設為0號學生
        enroll = Enroll(classroom_id=self.object.id, student_id=self.request.user.id, seat=0)
        enroll.save()   
        return redirect("/teacher/classroom")   
			
    def get_context_data(self, **kwargs):
        context = super(ClassroomCreateView, self).get_context_data(**kwargs)
        context['domains'] = Domain.objects.all()
        context['levels'] = Level.objects.all()
        return context	
        
# 修改選課密碼
def classroom_edit(request, classroom_id):
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id):
        return redirect("homepage")
    classroom = Classroom.objects.get(id=classroom_id)
    if request.method == 'POST':
        form = ClassroomForm(request.POST)
        if form.is_valid():
            classroom.name =form.cleaned_data['name']
            classroom.password = form.cleaned_data['password']
            classroom.domains = request.POST.getlist('domains')
            classroom.levels = request.POST.getlist('levels')	
            classroom.save()
            return redirect('/teacher/classroom')
    else:
        form = ClassroomForm(instance=classroom)
        domains = Domain.objects.all()
        levels = Level.objects.all()
    return render(request,'teacher/classroom_form.html',{'form': form, 'classroom': classroom, 'domains':domains, 'levels':levels})        
    
# 設定班級助教
def classroom_assistant(request, classroom_id):
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id):
        return redirect("homepage")
    assistants = Assistant.objects.filter(classroom_id=classroom_id).order_by("-id")
    classroom = Classroom.objects.get(id=classroom_id)

    return render(request,'teacher/assistant.html',{'assistants': assistants, 'classroom':classroom})        

# 教師可以查看所有帳號
class AssistantListView(ListView):
    context_object_name = 'users'
    paginate_by = 20
    template_name = 'teacher/assistant_user.html'
    
    def get_queryset(self):        
        if self.request.GET.get('account') != None:
            keyword = self.request.GET.get('account')
            queryset = User.objects.filter(Q(username__icontains=keyword) | Q(first_name__icontains=keyword)).order_by('-id')
        else :
            queryset = User.objects.all().order_by('-id')				
        return queryset

    def get_context_data(self, **kwargs):
        context = super(AssistantListView, self).get_context_data(**kwargs)
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        assistant_list = []
        assistants = Assistant.objects.filter(classroom_id=self.kwargs['classroom_id'])
        for assistant in assistants:
            assistant_list.append(assistant.user_id)
        context['assistants'] = assistant_list
        return context	

# 列出所有助教課程
class AssistantClassroomListView(ListView):
    model = Classroom
    context_object_name = 'classrooms'
    template_name = 'teacher/assistant_list.html'
    paginate_by = 20
    def get_queryset(self):      
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        classroom_list = []
        for assistant in assistants:
            classroom_list.append(assistant.classroom_id)
        queryset = Classroom.objects.filter(id__in=classroom_list).order_by("-id")
        return queryset
            
# Ajax 設為助教、取消助教
def assistant_make(request):
    classroom_id = request.POST.get('classroomid')	
    user_id = request.POST.get('userid')
    action = request.POST.get('action')
    if user_id and action :
        if action == 'set':            
            try :
                assistant = Assistant.objects.get(classroom_id=classroom_id, user_id=user_id) 	
            except ObjectDoesNotExist :
                assistant = Assistant(classroom_id=classroom_id, user_id=user_id) 
                assistant.save()
            # 將教師設為0號學生
            enroll = Enroll(classroom_id=classroom_id, student_id=user_id, seat=0)
            enroll.save() 
        else : 
            try :
                assistant = Assistant.objects.get(classroom_id=classroom_id, user_id=user_id)
                assistant.delete()
                enroll = Enroll.objects.filter(classroom_id=classroom_id, student_id=user_id)
                enroll.delete()								
            except ObjectDoesNotExist :
                pass             
        return JsonResponse({'status':'ok'}, safe=False)
    else:
        return JsonResponse({'status':'fail'}, safe=False)
	
# 退選
def unenroll(request, enroll_id, classroom_id):
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id):
        return redirect("homepage")    
    enroll = Enroll.objects.get(id=enroll_id)
    enroll.delete()
    classroom_name = Classroom.objects.get(id=classroom_id).name
    return redirect('/student/classmate/'+classroom_id)  


# 列出所有討論主題
class ForumListView(ListView):
    model = FWork
    context_object_name = 'forums'
    template_name = "teacher/forum_list.html"		
    paginate_by = 20
    def get_queryset(self):        
        fclasses = FClass.objects.filter(classroom_id=self.kwargs['classroom_id']).order_by("-publication_date", "-forum_id")
        forums = []
        for fclass in fclasses:
            forum = FWork.objects.get(id=fclass.forum_id)
            forums.append([forum, fclass])
        return forums
			
    def get_context_data(self, **kwargs):
        context = super(ForumListView, self).get_context_data(**kwargs)
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['classroom'] = classroom
        return context	
        
#新增一個討論主題
class ForumCreateView(CreateView):
    model = FWork
    form_class = ForumForm
    template_name = "teacher/forum_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.teacher_id = self.request.user.id
        self.object.classroom_id = self.kwargs['classroom_id']
        self.object.domains = self.request.POST.getlist('domains')
        self.object.levels = self.request.POST.getlist('levels')	        
        self.object.save()  
        classrooms = self.request.POST.getlist('classrooms')
        for classroom in classrooms:
          forum_class = FClass(forum_id=self.object.id, classroom_id=classroom)
          forum_class.save()
        
        return redirect("/teacher/forum/"+self.kwargs['classroom_id'])           
        
    def get_context_data(self, **kwargs):
        context = super(ForumCreateView, self).get_context_data(**kwargs)
        classroom_list = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id)
        for classroom in classrooms:
            classroom_list.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            if not assistant.classroom_id in classroom_list:
                classroom_list.append(assistant.classroom_id)
        classrooms = Classroom.objects.filter(id__in=classroom_list).order_by("-id")
        context['classrooms'] = classrooms
        context['classroom_id'] = int(self.kwargs['classroom_id'])
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['domains'] = Domain.objects.all()
        context['levels'] = Level.objects.all()
        return context	
  
        return redirect("/teacher/forum/"+self.kwargs['classroom_id'])        
	
def forum_categroy(request, classroom_id, forum_id):
    forum = FWork.objects.get(id=forum_id)
    domains = Domain.objects.all()
    levels = Level.objects.all()		
    if request.method == 'POST':
        form = ForumCategroyForm(request.POST)
        if form.is_valid():
            forum.domains = request.POST.getlist('domains')
            forum.levels = request.POST.getlist('levels')	
            forum.save()
            return redirect('/teacher/forum/'+classroom_id+'/#'+str(forum.id))
    else:
        form = ForumCategroyForm(instance=forum)
        
    return render(request,'teacher/forum_categroy_form.html',{'domains': domains, 'levels':levels, 'classroom_id': classroom_id, 'forum':forum})

	
# 列出所有討論主題
class ForumAllListView(ListView):
    model = FWork
    context_object_name = 'forums'
    template_name = "teacher/forum_all.html"		
    paginate_by = 20
		
    def get_queryset(self):
      # 年級
      if self.kwargs['categroy'] == "1":
        queryset = FWork.objects.filter(levels__contains=self.kwargs['categroy_id']).order_by("-id")
      # 學習領域
      elif self.kwargs['categroy'] == "2":
        queryset = FWork.objects.filter(domains__contains=self.kwargs['categroy_id']).order_by("-id")   
      else:
        queryset = FWork.objects.all().order_by("-id")
      if self.request.GET.get('account') != None:
        keyword = self.request.GET.get('account')
        users = User.objects.filter(Q(username__icontains=keyword) | Q(first_name__icontains=keyword)).order_by('-id')
        user_list = []
        for user in users:
            user_list.append(user.id)
        forums = queryset.filter(teacher_id__in=user_list)
        return forums
      else:				
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(ForumAllListView, self).get_context_data(**kwargs)
        context['categroy'] = self.kwargs['categroy']							
        context['categroy_id'] = self.kwargs['categroy_id']							
        context['levels'] = Level.objects.all()				
        context['domains'] = Domain.objects.all()
        return context	

# 展示討論素材
def forum_show(request, forum_id):
    forum = FWork.objects.get(id=forum_id)
    domains = Domain.objects.all()
    domain_dict = {}
    for domain in domains :
        key = domain.id
        domain_dict[key] = domain
    levels = Level.objects.all()	
    level_dict = {}
    for level in levels :
        key = level.id
        level_dict[key] = level
    contents = FContent.objects.filter(forum_id=forum_id)
    domains = []		
    if forum.domains:
        forum_domains = ast.literal_eval(forum.domains)
        for domain in forum_domains:
            key = int(domain)
            domains.append(domain_dict[key])
    levels = []						
    if forum.levels:
        forum_levels = ast.literal_eval(forum.levels)
        for level in forum_levels:
            key = int(level)			
            levels.append(level_dict[key])
    return render(request,'teacher/forum_show.html',{'domains':domains, 'levels':levels, 'contents':contents, 'forum':forum})

		
# 列出某討論主題的班級
class ForumClassListView(ListView):
    model = FWork
    context_object_name = 'classrooms'
    template_name = "teacher/forum_class.html"		
    paginate_by = 20
	
    def get_queryset(self):        		
        fclass_dict = dict(((fclass.classroom_id, fclass) for fclass in FClass.objects.filter(forum_id=self.kwargs['forum_id'])))		
        classroom_list = []
        classroom_ids = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id).order_by("-id")
        for classroom in classrooms:
            if classroom.id in fclass_dict:
                classroom_list.append([classroom, True, fclass_dict[classroom.id].deadline, fclass_dict[classroom.id].deadline_date])
            else :
                classroom_list.append([classroom, False, False, timezone.now()])
            classroom_ids.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            classroom = Classroom.objects.get(id=assistant.classroom_id)
            if not classroom.id in classroom_ids:
                if classroom.id in fclass_dict:
                    classroom_list.append([classroom, True, fclass_dict[classroom.id].deadline, fclass_dict[classroom.id].deadline_date])
                else :
                    classroom_list.append([classroom, False, False, timezone.now()])
        return classroom_list
			
    def get_context_data(self, **kwargs):
        context = super(ForumClassListView, self).get_context_data(**kwargs)				
        fwork = FWork.objects.get(id=self.kwargs['forum_id'])
        context['fwork'] = fwork
        context['forum_id'] = self.kwargs['forum_id']
        return context	
	
# Ajax 開放班取、關閉班級
def forum_switch(request):
    forum_id = request.POST.get('forumid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        fwork = FClass.objects.get(forum_id=forum_id, classroom_id=classroom_id)
        if status == 'false' :
    				fwork.delete()
    except ObjectDoesNotExist:
        if status == 'true':
            fwork = FClass(forum_id=forum_id, classroom_id=classroom_id)
            fwork.save()
    return JsonResponse({'status':status}, safe=False)        
	
# 列出某作業所有同學名單
def forum_class(request, classroom_id, work_id):
    enrolls = Enroll.objects.filter(classroom_id=classroom_id)
    classroom_name = Classroom.objects.get(id=classroom_id).name
    classmate_work = []
    scorer_name = ""
    for enroll in enrolls:
        try:    
            work = SWork.objects.get(student_id=enroll.student_id, index=work_id)
            if work.scorer > 0 :
                scorer = User.objects.get(id=work.scorer)
                scorer_name = scorer.first_name
            else :
                scorer_name = "1"
        except ObjectDoesNotExist:
            work = SWork(index=work_id, student_id=1)
        try:
            group_name = EnrollGroup.objects.get(id=enroll.group).name
        except ObjectDoesNotExist:
            group_name = "沒有組別"
        assistant = Assistant.objects.filter(classroom_id=classroom_id, student_id=enroll.student_id, lesson=work_id)
        if assistant.exists():
            classmate_work.append([enroll,work,1, scorer_name, group_name])
        else :
            classmate_work.append([enroll,work,0, scorer_name, group_name])   
    def getKey(custom):
        return custom[0].seat
	
    classmate_work = sorted(classmate_work, key=getKey)
   
    return render(request,'teacher/twork_class.html',{'classmate_work': classmate_work, 'classroom_id':classroom_id, 'index': work_id})

# 列出所有討論主題素材
class ForumContentListView(ListView):
    model = FContent
    context_object_name = 'contents'
    template_name = "teacher/forum_content.html"		
    def get_queryset(self):
        queryset = FContent.objects.filter(forum_id=self.kwargs['forum_id']).order_by("-id")
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(ForumContentListView, self).get_context_data(**kwargs)
        fwork = FWork.objects.get(id=self.kwargs['forum_id'])
        fclasses = FClass.objects.filter(forum_id=self.kwargs['forum_id'])				
        context['fwork']= fwork
        context['forum_id'] = self.kwargs['forum_id']
        context['fclasses'] = fclasses
        return context	
			
#新增一個課程
class ForumContentCreateView(CreateView):
    model = FContent
    form_class = ForumContentForm
    template_name = "teacher/forum_content_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        work = FContent(forum_id=self.object.forum_id)
        if self.object.types == 1:
            work.types = 1
            work.title = self.object.title
            work.link = self.object.link
        if self.object.types  == 2:
            work.types = 2					
            work.youtube = self.object.youtube
        if self.object.types  == 3:
            work.types = 3
            myfile = self.request.FILES['content_file']
            fs = FileSystemStorage()
            filename = uuid4().hex
            work.title = myfile.name
            work.filename = str(self.request.user.id)+"/"+filename
            fs.save("static/upload/"+str(self.request.user.id)+"/"+filename, myfile)
        if self.object.types  == 4:
            work.types = 4
        work.memo = self.object.memo
        work.save()         
  
        return redirect("/teacher/forum/content/"+self.kwargs['forum_id'])  

    def get_context_data(self, **kwargs):
        ctx = super(ForumContentCreateView, self).get_context_data(**kwargs)
        ctx['forum'] = FWork.objects.get(id=self.kwargs['forum_id'])
        return ctx

def forum_delete(request, forum_id, content_id):
    instance = FContent.objects.get(id=content_id)
    instance.delete()

    return redirect("/teacher/forum/content/"+forum_id)  
	
def forum_edit(request, forum_id, content_id):
    try:
        instance = FContent.objects.get(id=content_id)
    except:
        pass
    if request.method == 'POST':
            content_id = request.POST.get("id", "")
            try:
                content = FContent.objects.get(id=content_id)
            except ObjectDoesNotExist:
	              content = FContent(forum_id= request.POST.get("forum_id", ""), types=form.cleaned_data['types'])
            if content.types == 1:
                content.title = request.POST.get("title", "")
                content.link = request.POST.get("link", "")
            elif content.types == 2:
                content.youtube = request.POST.get("youtube", "")
            elif content.types == 3:
                myfile =  request.FILES.get("content_file", "")
                fs = FileSystemStorage()
                filename = uuid4().hex
                content.title = myfile.name
                content.filename = str(request.user.id)+"/"+filename
                fs.save("static/upload/"+str(request.user.id)+"/"+filename, myfile)
            content.memo = request.POST.get("memo", "")
            content.save()
            return redirect('/teacher/forum/content/'+forum_id)   
    return render(request,'teacher/forum_edit.html',{'content': instance, 'forum_id':forum_id, 'content_id':content_id})		
	
def forum_download(request, content_id):
    content = FContent.objects.get(id=content_id)
    filename = content.title
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))		
    download =  BASE_DIR + "/static/upload/" + content.filename
    wrapper = FileWrapper(file( download, "r" ))
    response = HttpResponse(wrapper, content_type = 'application/force-download')
    #response = HttpResponse(content_type='application/force-download')
    response['Content-Disposition'] = 'attachment; filename={0}'.format(filename.encode('utf8'))
    # It's usually a good idea to set the 'Content-Length' header too.
    # You can also set any other required headers: Cache-Control, etc.
    return response
    #return render(request,'student/download.html', {'download':download})
		
class ForumEditUpdateView(UpdateView):
    model = FWork
    fields = ['title']
    template_name = 'form.html'
    #success_url = '/teacher/forum/domain/'
    def get_success_url(self):
        succ_url =  '/teacher/forum/'+self.kwargs['classroom_id']
        return succ_url
	
def forum_export(request, classroom_id, forum_id):
	if not is_teacher(request.user, classroom_id):
		return redirect("/")
	classroom = Classroom.objects.get(id=classroom_id)
	try:
		fwork = FWork.objects.get(id=forum_id)
		enrolls = Enroll.objects.filter(classroom_id=classroom_id)
		datas = []
		contents = FContent.objects.filter(forum_id=forum_id).order_by("-id")
		fwork = FWork.objects.get(id=forum_id)
		works_pool = SFWork.objects.filter(index=forum_id).order_by("-id")
		reply_pool = SFReply.objects.filter(index=forum_id).order_by("-id")	
		file_pool = SFContent.objects.filter(index=forum_id, visible=True).order_by("-id")	
		for enroll in enrolls:
			works = filter(lambda w: w.student_id==enroll.student_id, works_pool)
			if len(works)>0:
				replys = filter(lambda w: w.work_id==works[0].id, reply_pool)
			else:
				replys = []
			files = filter(lambda w: w.student_id==enroll.student_id, file_pool)
			if enroll.seat > 0:
				datas.append([enroll, works, replys, files])
		def getKey(custom):
			return -custom[0].seat
		datas = sorted(datas, key=getKey, reverse=True)	
		#word
		document = Document()
		docx_title=u"討論區-" + classroom.name + "-"+ str(timezone.localtime(timezone.now()).date())+".docx"
		document.add_paragraph(request.user.first_name + u'的討論區作業')
		document.add_paragraph(u'主題：'+fwork.title)		
		document.add_paragraph(u"班級：" + classroom.name)		
		
		for enroll, works, replys, files in datas:
			user = User.objects.get(id=enroll.student_id)
			run = document.add_paragraph().add_run(str(enroll.seat)+")"+user.first_name)
			font = run.font
			font.color.rgb = RGBColor(0xFA, 0x24, 0x00)
			if len(works)>0:
				#p = document.add_paragraph(str(works[0].publication_date)[:19]+'\n'+works[0].memo)
				p = document.add_paragraph(str(localtime(works[0].publication_date))[:19]+'\n')
				# 將 memo 以時間標記為切割點，切分為一堆 tokens
				tokens = re.split('(\[m_\d+#\d+:\d+:\d+\])', works[0].memo)
				# 依續比對 token 格式
				for token in tokens:
					m = re.match('\[m_(\d+)#(\d+):(\d+):(\d+)\]', token)
					if m: # 若為時間標記，則插入連結
						vid = filter(lambda material: material.id == int(m.group(1)), contents)[0]
						add_hyperlink(document, p, vid.youtube+"&t="+m.group(2)+"h"+m.group(3)+"m"+m.group(4)+"s", "["+m.group(2)+":"+m.group(3)+":"+m.group(4)+"]")
					else: # 以一般文字插入
						p.add_run(token)
			if len(replys)>0:
				for reply in replys:
					user = User.objects.get(id=reply.user_id)
					run = document.add_paragraph().add_run(user.first_name+u'>'+str(localtime(reply.publication_date))[:19]+u'>留言:\n'+reply.memo)
					font = run.font
					font.color.rgb = RGBColor(0x42, 0x24, 0xE9)		
			if len(files)>0:
				for file in files:
					if file.visible:
						if file.title[-3:].upper() == "PNG" or file.title[-3:].upper() == "JPG":
							filename = 'static/upload/'+file.filename
							if os.path.exists(filename):
								copyfile(filename, 'static/upload/file.png')					
								document.add_picture('static/upload/file.png',width=Inches(6.0))
						else:
							p = document.add_paragraph()
							full_url = request.build_absolute_uri()
							index = full_url.find("/",9)
							url = full_url[:index] + "/student/forum/download/" + str(file.id) 
							add_hyperlink(document, p, url, file.title)
		# Prepare document for download        
		f = StringIO.StringIO()
		document.save(f)
		length = f.tell()
		f.seek(0)
		response = HttpResponse(
			f.getvalue(),
			content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
		)
		response['Content-Disposition'] = 'attachment; filename={0}'.format(docx_title.encode('utf8')) 
		response['Content-Length'] = length
		return response

	except ObjectDoesNotExist:
		pass
	return True

def add_hyperlink(document, paragraph, url, name):
    """
    Add a hyperlink to a paragraph.
    :param document: The Document being edited.
    :param paragraph: The Paragraph the hyperlink is being added to.
    :param url: The url to be added to the link.
    :param name: The text for the link to be displayed in the paragraph
    :return: None
    """

    part = document.part
    rId = part.relate_to(url, RT.HYPERLINK, is_external=True)

    init_hyper = OxmlElement('w:hyperlink')
    init_hyper.set(qn('r:id'), rId, )
    init_hyper.set(qn('w:history'), '1')

    new_run = OxmlElement('w:r')

    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')

    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = name
    init_hyper.append(new_run)

    r = paragraph.add_run()
    r._r.append(init_hyper)
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return None

def forum_grade(request, classroom_id, action):
	classroom = Classroom.objects.get(id=classroom_id)
	forum_ids = []
	forums = []
	fclasses = FClass.objects.filter(classroom_id=classroom_id).order_by("publication_date", "forum_id")
	for fclass in fclasses:
		forum_ids.append(fclass.forum_id)
		forum = FWork.objects.get(id=fclass.forum_id)
		forums.append(forum.title)
	enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")
	datas = {}
	for enroll in enrolls:
			sfworks = SFWork.objects.filter(index__in=forum_ids, student_id=enroll.student_id).order_by("id")
			if len(sfworks) > 0:
				for fclass in fclasses:
						works = filter(lambda w: w.index==fclass.forum_id, sfworks)
						if enroll.student_id in datas:
							if len(works) > 0 :
								datas[enroll.student_id].append(works[0])
							else :
								datas[enroll.student_id].append(SFWork())
						else:
							if len(works) > 0:
								datas[enroll.student_id] = [works[0]]
							else :
								datas[enroll.student_id] = [SFWork()]
			else :
				datas[enroll.student_id] = [SFWork()]
	results = []
	for enroll in enrolls:
		student_name = User.objects.get(id=enroll.student_id).first_name
		results.append([enroll, student_name, datas[enroll.student_id]])
	
	#下載Excel
	if action == "1":
		classroom = Classroom.objects.get(id=classroom_id)       
		output = StringIO.StringIO()
		workbook = xlsxwriter.Workbook(output)    
		worksheet = workbook.add_worksheet(classroom.name)
		date_format = workbook.add_format({'num_format': 'yy/mm/dd'})
		
		row = 1
		worksheet.write(row, 1, u'座號')
		worksheet.write(row, 2, u'姓名')
		index = 3
		for forum in forums:
			worksheet.write(row, index, forum)
			index += 1
		
		row += 1
		index = 3
		for fclass in fclasses:
			worksheet.write(row, index, datetime.strptime(str(fclass.publication_date)[:19],'%Y-%m-%d %H:%M:%S'), date_format)
			index += 1			

		for enroll, student_name, works in results:
			row += 1
			worksheet.write(row, 1, enroll.seat)
			worksheet.write(row, 2, student_name)
			index = 3
			for work in works:
				if work.id:
					worksheet.write(row, index, work.score)
				else:
					worksheet.write(row, index, '')
				index +=1 

		workbook.close()
		# xlsx_data contains the Excel file
		response = HttpResponse(content_type='application/vnd.ms-excel')
		filename = classroom.name + '-' + str(localtime(timezone.now()).date()) + '.xlsx'
		response['Content-Disposition'] = 'attachment; filename={0}'.format(filename.encode('utf8'))
		xlsx_data = output.getvalue()
		response.write(xlsx_data)
		return response
	else :
		return render(request,'teacher/forum_grade.html',{'results':results, 'forums':forums, 'classroom_id':classroom_id, 'fclasses':fclasses})

def forum_deadline(request, classroom_id, forum_id):
    forum = FWork.objects.get(id=forum_id)
    classroom = Classroom.objects.get(id=classroom_id)
    if request.method == 'POST':
        form = CategroyForm(request.POST)
        if form.is_valid():
            forum.domains = request.POST.getlist('domains')
            forum.levels = request.POST.getlist('levels')	
            forum.save()
            return redirect('/teacher/forum/'+classroom_id)
    else:
        fclass = FClass.objects.get(classroom_id=classroom_id, forum_id=forum_id)
        form = ForumDeadlineForm(instance=fclass)
        fclasses = FClass.objects.filter(forum_id=forum_id).order_by("-id")
    return render(request,'teacher/forum_deadline_form.html',{'fclasses':fclasses, 'fclass':fclass, 'forum':forum, 'classroom':classroom})

	
# Ajax 設定期限、取消期限
def forum_deadline_set(request):
    forum_id = request.POST.get('forumid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        fclass = FClass.objects.get(forum_id=forum_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        fclass = Fclass(forum_id=forum_id, classroom_id=classroom_id)
    if status == 'True':
        fclass.deadline = True
    else :
        fclass.deadline = False
    fclass.save()
    return JsonResponse({'status':status}, safe=False)        

# Ajax 設定期限日期
def forum_deadline_date(request):
    forum_id = request.POST.get('forumid')
    classroom_id = request.POST.get('classroomid')		
    deadline_date = request.POST.get('deadlinedate')
    try:
        fclass = FClass.objects.get(forum_id=forum_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        fclass = FClass(forum_id=forum_id, classroom_id=classroom_id)
    #fclass.deadline_date = deadline_date.strftime('%d/%m/%Y')
    fclass.deadline_date = datetime.strptime(deadline_date, '%Y %B %d - %H:%M')
    fclass.save()
    return JsonResponse({'status':deadline_date}, safe=False)             
        
#新增一個公告
class AnnounceCreateView(CreateView):
    model = Message
    form_class = AnnounceForm
    template_name = 'teacher/announce_form.html'     
    def form_valid(self, form):
        self.object = form.save(commit=False)			
        classrooms = self.request.POST.getlist('classrooms')
        files = []
        if self.request.FILES.getlist('files'):
             for file in self.request.FILES.getlist('files'):
                fs = FileSystemStorage()
                filename = uuid4().hex							
                fs.save("static/upload/"+str(self.request.user.id)+"/"+filename, file)								
                files.append([filename, file.name])		
        for classroom_id in classrooms:
            message = Message()
            message.title = u"[公告]" + self.request.user.first_name + ":" + self.object.title
            message.author_id = self.request.user.id	
            message.type = 1 #公告
            message.classroom_id = classroom_id
            message.content = self.object.content
            message.save()
            message.url = "/account/line/detail/" + classroom_id + "/" + str(message.id)
            message.save()
            if files:
                for file, name in files:
                    content = MessageContent()
                    content.title = name
                    content.message_id = message.id
                    content.filename = str(self.request.user.id)+"/"+file
                    content.save()		

            # 班級學生訊息
            enrolls = Enroll.objects.filter(classroom_id=classroom_id)
            for enroll in enrolls:
                messagepoll = MessagePoll(message_type=1, message_id=message.id, reader_id=enroll.student_id, classroom_id=classroom_id)
                messagepoll.save()               
        return redirect("/student/announce/"+self.kwargs['classroom_id']) 
			
    def get_context_data(self, **kwargs):
        context = super(AnnounceCreateView, self).get_context_data(**kwargs)
        context['class'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        classroom_list = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id)
        for classroom in classrooms:
            classroom_list.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            if not assistant.classroom_id in classroom_list:
                classroom_list.append(assistant.classroom_id)
        classrooms = Classroom.objects.filter(id__in=classroom_list).order_by("-id")
        context['classrooms'] = classrooms
        return context	   
        
    # 限本班任課教師        
    def render_to_response(request,self, context):
        if not is_teacher(self.request.user, self.kwargs['classroom_id']):
            if not is_assistant(self.request.user, self.kwargs['classroom_id']) :
                return redirect('/')
        return super(AnnounceCreateView, self).render(request,context)        
			
'''    
----------------------- 思辨區
'''
# 列出所有思辨主題
class SpeculationListView(ListView):
    model = SpeculationWork
    context_object_name = 'forums'
    template_name = "teacher/speculation_list.html"		
    paginate_by = 20
    def get_queryset(self):        
        fclasses = SpeculationClass.objects.filter(classroom_id=self.kwargs['classroom_id']).order_by("-publication_date", "-forum_id")
        forums = []
        for fclass in fclasses:
            forum = SpeculationWork.objects.get(id=fclass.forum_id)
            forums.append([forum, fclass])
        return forums
			
    def get_context_data(self, **kwargs):
        context = super(SpeculationListView, self).get_context_data(**kwargs)
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['classroom'] = classroom
        return context	
        
#新增一個討論主題
class SpeculationCreateView(CreateView):
    model = SpeculationWork
    form_class = SpeculationForm
    template_name = "teacher/speculation_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.teacher_id = self.request.user.id
        self.object.classroom_id = self.kwargs['classroom_id']
        self.object.domains = self.request.POST.getlist('domains')
        self.object.levels = self.request.POST.getlist('levels')	        
        self.object.save()  
        classrooms = self.request.POST.getlist('classrooms')
        for classroom in classrooms:
          forum_class = SpeculationClass(forum_id=self.object.id, classroom_id=classroom)
          forum_class.save()
        
        return redirect("/teacher/speculation/"+self.kwargs['classroom_id'])           
        
    def get_context_data(self, **kwargs):
        context = super(SpeculationCreateView, self).get_context_data(**kwargs)
        classroom_list = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id)
        for classroom in classrooms:
            classroom_list.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            if not assistant.classroom_id in classroom_list:
                classroom_list.append(assistant.classroom_id)
        classrooms = Classroom.objects.filter(id__in=classroom_list).order_by("-id")
        context['classrooms'] = classrooms
        context['classroom_id'] = int(self.kwargs['classroom_id'])
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['domains'] = Domain.objects.all()
        context['levels'] = Level.objects.all()
        return context	
  
        return redirect("/teacher/speculation/"+self.kwargs['classroom_id'])        
	
def speculation_categroy(request, classroom_id, forum_id):
    forum = SpeculationWork.objects.get(id=forum_id)
    domains = Domain.objects.all()
    levels = Level.objects.all()		
    if request.method == 'POST':
        form = SpeculationCategroyForm(request.POST)
        if form.is_valid():
            forum.domains = request.POST.getlist('domains')
            forum.levels = request.POST.getlist('levels')	
            forum.save()
            return redirect('/teacher/speculation/'+classroom_id+'/#'+str(forum.id))
    else:
        form = SpeculationCategroyForm(instance=forum)
        
    return render(request,'teacher/speculation_categroy_form.html',{'domains': domains, 'levels':levels, 'classroom_id': classroom_id, 'speculation':forum})

	
# 列出所有思辨主題
class SpeculationAllListView(ListView):
    model = SpeculationWork
    context_object_name = 'forums'
    template_name = "teacher/speculation_all.html"		
    paginate_by = 20
		
    def get_queryset(self):
      # 年級
      if self.kwargs['categroy'] == "1":
        queryset = SpeculationWork.objects.filter(levels__contains=self.kwargs['categroy_id']).order_by("-id")
      # 學習領域
      elif self.kwargs['categroy'] == "2":
        queryset = SpeculationWork.objects.filter(domains__contains=self.kwargs['categroy_id']).order_by("-id")   
      else:
        queryset = SpeculationWork.objects.all().order_by("-id")
      if self.request.GET.get('account') != None:
        keyword = self.request.GET.get('account')
        users = User.objects.filter(Q(username__icontains=keyword) | Q(first_name__icontains=keyword)).order_by('-id')
        user_list = []
        for user in users:
            user_list.append(user.id)
        forums = queryset.filter(teacher_id__in=user_list)
        return forums
      else:				
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(SpeculationAllListView, self).get_context_data(**kwargs)
        context['categroy'] = self.kwargs['categroy']							
        context['categroy_id'] = self.kwargs['categroy_id']							
        context['levels'] = Level.objects.all()				
        context['domains'] = Domain.objects.all()
        return context	

# 展示思辨素材
def speculation_show(request, forum_id):
    forum = SpeculationWork.objects.get(id=forum_id)
    domains = Domain.objects.all()
    domain_dict = {}
    for domain in domains :
        key = domain.id
        domain_dict[key] = domain
    levels = Level.objects.all()	
    level_dict = {}
    for level in levels :
        key = level.id
        level_dict[key] = level
    contents = SpeculationContent.objects.filter(forum_id=forum_id)
    domains = []		
    if forum.domains:
        forum_domains = ast.literal_eval(forum.domains)
        for domain in forum_domains:
            key = int(domain)
            domains.append(domain_dict[key])
    levels = []						
    if forum.levels:
        forum_levels = ast.literal_eval(forum.levels)
        for level in forum_levels:
            key = int(level)			
            levels.append(level_dict[key])
    return render(request,'teacher/speculation_show.html',{'domains':domains, 'levels':levels, 'contents':contents, 'forum':forum})

		
# 列出某思辨主題的班級
class SpeculationClassListView(ListView):
    model = SpeculationWork
    context_object_name = 'classrooms'
    template_name = "teacher/speculation_class.html"		
    paginate_by = 20
	
    def get_queryset(self):        		
        fclass_dict = dict(((fclass.classroom_id, fclass) for fclass in SpeculationClass.objects.filter(forum_id=self.kwargs['forum_id'])))		
        classroom_list = []
        classroom_ids = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id).order_by("-id")
        for classroom in classrooms:
            if classroom.id in fclass_dict:
                classroom_list.append([classroom, True, fclass_dict[classroom.id].deadline, fclass_dict[classroom.id].deadline_date])
            else :
                classroom_list.append([classroom, False, False, timezone.now()])
            classroom_ids.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            classroom = Classroom.objects.get(id=assistant.classroom_id)
            if not classroom.id in classroom_ids:
                if classroom.id in fclass_dict:
                    classroom_list.append([classroom, True, fclass_dict[classroom.id].deadline, fclass_dict[classroom.id].deadline_date])
                else :
                    classroom_list.append([classroom, False, False, timezone.now()])
        return classroom_list
			
    def get_context_data(self, **kwargs):
        context = super(SpeculationClassListView, self).get_context_data(**kwargs)				
        fwork = SpeculationWork.objects.get(id=self.kwargs['forum_id'])
        context['fwork'] = fwork
        context['forum_id'] = self.kwargs['forum_id']
        return context	
	
# Ajax 開放班取、關閉班級
def speculation_switch(request):
    forum_id = request.POST.get('forumid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        fwork = SpeculationClass.objects.get(forum_id=forum_id, classroom_id=classroom_id)
        if status == 'false' :
            fwork.delete()
    except ObjectDoesNotExist:
        if status == 'true':
            fwork = SpeculationClass(forum_id=forum_id, classroom_id=classroom_id)
            fwork.save()
    return JsonResponse({'status':status}, safe=False)        
	
# 列出某作業所有同學名單
def speculation_class(request, classroom_id, work_id):
    enrolls = Enroll.objects.filter(classroom_id=classroom_id)
    classroom_name = Classroom.objects.get(id=classroom_id).name
    classmate_work = []
    scorer_name = ""
    for enroll in enrolls:
        try:    
            work = SWork.objects.get(student_id=enroll.student_id, index=work_id)
            if work.scorer > 0 :
                scorer = User.objects.get(id=work.scorer)
                scorer_name = scorer.first_name
            else :
                scorer_name = "1"
        except ObjectDoesNotExist:
            work = SWork(index=work_id, student_id=1)
        try:
            group_name = EnrollGroup.objects.get(id=enroll.group).name
        except ObjectDoesNotExist:
            group_name = "沒有組別"
        assistant = Assistant.objects.filter(classroom_id=classroom_id, student_id=enroll.student_id, lesson=work_id)
        if assistant.exists():
            classmate_work.append([enroll,work,1, scorer_name, group_name])
        else :
            classmate_work.append([enroll,work,0, scorer_name, group_name])   
    def getKey(custom):
        return custom[0].seat
	
    classmate_work = sorted(classmate_work, key=getKey)
   
    return render(request,'teacher/speculation_class.html',{'classmate_work': classmate_work, 'classroom_id':classroom_id, 'index': work_id})

# 列出所有思辨主題素材
class SpeculationContentListView(ListView):
    model = SpeculationContent
    context_object_name = 'contents'
    template_name = "teacher/speculation_content.html"		
    def get_queryset(self):
        queryset = SpeculationContent.objects.filter(forum_id=self.kwargs['forum_id'])
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(SpeculationContentListView, self).get_context_data(**kwargs)
        fwork = SpeculationWork.objects.get(id=self.kwargs['forum_id'])
        context['fwork']= fwork
        context['forum_id'] = self.kwargs['forum_id']
        return context	
			
#新增一個課程
class SpeculationContentCreateView(CreateView):
    model = SpeculationContent
    form_class = SpeculationContentForm
    template_name = "teacher/speculation_content_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        work = SpeculationContent(forum_id=self.object.forum_id)
        if self.object.types == 1:
            work.types = 1
            work.text = self.object.text
        elif self.object.types  == 2:
            work.types = 2					
            work.youtube = self.object.youtube
        elif self.object.types  == 3:
            work.types = 3
            myfile = self.request.FILES['content_file']
            fs = FileSystemStorage()
            filename = uuid4().hex
            work.title = myfile.name
            work.filename = str(self.request.user.id)+"/"+filename
            fs.save("static/upload/"+str(self.request.user.id)+"/"+filename, myfile)	
        elif self.object.types  == 4:
            work.types = 4
            work.link = self.object.link
            work.title = self.object.title
        work.memo = self.object.memo
        work.save()         

        if self.object.types == 3:
          return JsonResponse({'files': [{'name': work.filename}]}, safe=False)
        return redirect("/teacher/speculation/content/"+self.kwargs['forum_id'])  

    def get_context_data(self, **kwargs):
        ctx = super(SpeculationContentCreateView, self).get_context_data(**kwargs)
        ctx['forum'] = SpeculationWork.objects.get(id=self.kwargs['forum_id'])
        return ctx

def speculation_delete(request, forum_id, content_id):
    instance = SpeculationContent.objects.get(id=content_id)
    instance.delete()

    return redirect("/teacher/speculation/content/"+forum_id)  
	
def speculation_edit(request, forum_id, content_id):
    try:
        instance = SpeculationContent.objects.get(id=content_id)
    except:
        pass
    if request.method == 'POST':
            content_id = request.POST.get("id", "")
            try:
                content = SpeculationContent.objects.get(id=content_id)
            except ObjectDoesNotExist:
	              content = SpeculattionContent(forum_id= request.POST.get("forum_id", ""), types=form.cleaned_data['types'])
            if content.types == 1:
                content.text = request.POST.get("text", "")
            elif content.types == 2:
                content.youtube = request.POST.get("youtube", "")
            elif content.types == 3:
                myfile =  request.FILES.get("content_file", "")
                fs = FileSystemStorage()
                filename = uuid4().hex
                content.title = myfile.name
                content.filename = str(request.user.id)+"/"+filename
                fs.save("static/upload/"+str(request.user.id)+"/"+filename, myfile)
            elif content.types == 4:
                content.title = request.POST.get("title", "")
                content.link = request.POST.get("link", "")								
            content.memo = request.POST.get("memo", "")
            content.save()
            return redirect('/teacher/speculation/content/'+forum_id)   
    return render(request,'teacher/speculation_edit.html',{'content': instance, 'forum_id':forum_id, 'content_id':content_id})		
	
def speculation_download(request, content_id):
    content = SpeculationContent.objects.get(id=content_id)
    filename = content.title
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    download =  BASE_DIR + "/static/upload/" + content.filename
    wrapper = FileWrapper(file( download, "r" ))
    response = HttpResponse(wrapper, content_type = 'application/force-download')
    #response = HttpResponse(content_type='application/force-download')
    response['Content-Disposition'] = 'attachment; filename={0}'.format(filename.encode('utf8'))
    # It's usually a good idea to set the 'Content-Length' header too.
    # You can also set any other required headers: Cache-Control, etc.
    return response
    #return render(request,'student/download.html', {'download':download})
		
class SpeculationEditUpdateView(UpdateView):
    model = SpeculationWork
    fields = ['title']
    template_name = 'form.html'
    #success_url = '/teacher/forum/domain/'
    def get_success_url(self):
        succ_url =  '/teacher/speculation/'+self.kwargs['classroom_id']
        return succ_url
	
def speculation_export(request, classroom_id, forum_id):
	if not is_teacher(request.user, classroom_id):
		return redirect("/")
	classroom = Classroom.objects.get(id=classroom_id)
	try:
		fwork = FWork.objects.get(id=forum_id)
		enrolls = Enroll.objects.filter(classroom_id=classroom_id)
		datas = []
		contents = FContent.objects.filter(forum_id=forum_id).order_by("-id")
		fwork = FWork.objects.get(id=forum_id)
		works_pool = SFWork.objects.filter(index=forum_id).order_by("-id")
		reply_pool = SFReply.objects.filter(index=forum_id).order_by("-id")	
		file_pool = SFContent.objects.filter(index=forum_id, visible=True).order_by("-id")	
		for enroll in enrolls:
			works = filter(lambda w: w.student_id==enroll.student_id, works_pool)
			if len(works)>0:
				replys = filter(lambda w: w.work_id==works[0].id, reply_pool)
			else:
				replys = []
			files = filter(lambda w: w.student_id==enroll.student_id, file_pool)
			if enroll.seat > 0:
				datas.append([enroll, works, replys, files])
		def getKey(custom):
			return -custom[0].seat
		datas = sorted(datas, key=getKey, reverse=True)	
		#word
		document = Document()
		docx_title=u"思辨區-" + classroom.name + "-"+ str(timezone.localtime(timezone.now()).date())+".docx"
		document.add_paragraph(request.user.first_name + u'的思辨區作業')
		document.add_paragraph(u'主題：'+fwork.title)		
		document.add_paragraph(u"班級：" + classroom.name)		
		
		for enroll, works, replys, files in datas:
			user = User.objects.get(id=enroll.student_id)
			run = document.add_paragraph().add_run(str(enroll.seat)+")"+user.first_name)
			font = run.font
			font.color.rgb = RGBColor(0xFA, 0x24, 0x00)
			if len(works)>0:
				#p = document.add_paragraph(str(works[0].publication_date)[:19]+'\n'+works[0].memo)
				p = document.add_paragraph(str(localtime(works[0].publication_date))[:19]+'\n')
				# 將 memo 以時間標記為切割點，切分為一堆 tokens
				tokens = re.split('(\[m_\d+#\d+:\d+:\d+\])', works[0].memo)
				# 依續比對 token 格式
				for token in tokens:
					m = re.match('\[m_(\d+)#(\d+):(\d+):(\d+)\]', token)
					if m: # 若為時間標記，則插入連結
						vid = filter(lambda material: material.id == int(m.group(1)), contents)[0]
						add_hyperlink(document, p, vid.youtube+"&t="+m.group(2)+"h"+m.group(3)+"m"+m.group(4)+"s", "["+m.group(2)+":"+m.group(3)+":"+m.group(4)+"]")
					else: # 以一般文字插入
						p.add_run(token)
			if len(replys)>0:
				for reply in replys:
					user = User.objects.get(id=reply.user_id)
					run = document.add_paragraph().add_run(user.first_name+u'>'+str(localtime(reply.publication_date))[:19]+u'>留言:\n'+reply.memo)
					font = run.font
					font.color.rgb = RGBColor(0x42, 0x24, 0xE9)		
			if len(files)>0:
				for file in files:
					if file.visible:
						if file.title[-3:].upper() == "PNG" or file.title[-3:].upper() == "JPG":
							filename = 'static/upload/'+file.filename
							if os.path.exists(filename):
								copyfile(filename, 'static/upload/file.png')					
								document.add_picture('static/upload/file.png',width=Inches(6.0))
						else:
							p = document.add_paragraph()
							full_url = request.build_absolute_uri()
							index = full_url.find("/",9)
							url = full_url[:index] + "/student/speculation/download/" + str(file.id) 
							add_hyperlink(document, p, url, file.title)
		# Prepare document for download        
		f = StringIO.StringIO()
		document.save(f)
		length = f.tell()
		f.seek(0)
		response = HttpResponse(
			f.getvalue(),
			content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
		)
		response['Content-Disposition'] = 'attachment; filename={0}'.format(docx_title.encode('utf8')) 
		response['Content-Length'] = length
		return response

	except ObjectDoesNotExist:
		pass
	return True

def speculation_grade(request, classroom_id, action):
	classroom = Classroom.objects.get(id=classroom_id)
	forum_ids = []
	forums = []
	fclasses = SpeculationClass.objects.filter(classroom_id=classroom_id).order_by("publication_date", "forum_id")
	for fclass in fclasses:
		forum_ids.append(fclass.forum_id)
		forum = SpeculationWork.objects.get(id=fclass.forum_id)
		forums.append(forum.title)
	enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")
	datas = {}
	for enroll in enrolls:
			sfworks = SSpeculationWork.objects.filter(index__in=forum_ids, student_id=enroll.student_id).order_by("id")
			if len(sfworks) > 0:
				for fclass in fclasses:
						works = filter(lambda w: w.index==fclass.forum_id, sfworks)
						if enroll.student_id in datas:
							if len(works) > 0 :
								datas[enroll.student_id].append(works[0])
							else :
								datas[enroll.student_id].append(SSpeculationWork())
						else:
							if len(works) > 0:
								datas[enroll.student_id] = [works[0]]
							else :
								datas[enroll.student_id] = [SSpeculationWork()]
			else :
				datas[enroll.student_id] = [SSpeculationWork()]
	results = []
	for enroll in enrolls:
		student_name = User.objects.get(id=enroll.student_id).first_name
		results.append([enroll, student_name, datas[enroll.student_id]])
	
	#下載Excel
	if action == "1":
		classroom = Classroom.objects.get(id=classroom_id)       
		output = StringIO.StringIO()
		workbook = xlsxwriter.Workbook(output)    
		worksheet = workbook.add_worksheet(classroom.name)
		date_format = workbook.add_format({'num_format': 'yy/mm/dd'})
		
		row = 1
		worksheet.write(row, 1, u'座號')
		worksheet.write(row, 2, u'姓名')
		index = 3
	
		for forum in forums:
			worksheet.write(row, index, forum)
			index += 1
		
		row += 1
		index = 3

		for fclass in fclasses:
			worksheet.write(row, index, datetime.strptime(str(fclass.publication_date)[:19],'%Y-%m-%d %H:%M:%S'), date_format)
			index += 1			

		for enroll, student_name, sworks in results:
			row += 1
			worksheet.write(row, 1, enroll.seat)
			worksheet.write(row, 2, student_name)
			index = 3
			for work in sworks:
				if work.id:
					worksheet.write(row, index, work.score)
				else:
					worksheet.write(row, index, '')
				index +=1 

		workbook.close()
		# xlsx_data contains the Excel file
		response = HttpResponse(content_type='application/vnd.ms-excel')
		filename = classroom.name + '-' + str(localtime(timezone.now()).date()) + '.xlsx'
		response['Content-Disposition'] = 'attachment; filename={0}'.format(filename.encode('utf8'))
		xlsx_data = output.getvalue()
		response.write(xlsx_data)
		return response
	else :
		return render(request,'teacher/speculation_grade.html',{'results':results, 'forums':forums, 'classroom_id':classroom_id, 'fclasses':fclasses})

def speculation_deadline(request, classroom_id, forum_id):
    forum = SpeculationWork.objects.get(id=forum_id)
    if request.method == 'POST':
        form = SpeculationCategroyForm(request.POST)
        if form.is_valid():
            forum.domains = request.POST.getlist('domains')
            forum.levels = request.POST.getlist('levels')	
            forum.save()
            return redirect('/teacher/speculation/'+classroom_id)
    else:
        fclass = SpeculationClass.objects.get(classroom_id=classroom_id, forum_id=forum_id)
        form = SpeculationDeadlineForm(instance=fclass)
    return render(request,'teacher/speculation_deadline_form.html',{'fclass':fclass})

	
# Ajax 設定期限、取消期限
def speculation_deadline_set(request):
    forum_id = request.POST.get('forumid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        fclass = SpeculationClass.objects.get(forum_id=forum_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        fclass = SpeculationClass(forum_id=forum_id, classroom_id=classroom_id)
    if status == 'True':
        fclass.deadline = True
    else :
        fclass.deadline = False
    fclass.save()
    return JsonResponse({'status':status}, safe=False)        

# Ajax 設定期限日期
def speculation_deadline_date(request):
    forum_id = request.POST.get('forumid')
    classroom_id = request.POST.get('classroomid')		
    deadline_date = request.POST.get('deadlinedate')
    try:
        fclass = SpeculationClass.objects.get(forum_id=forum_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        fclass = SpeculationClass(forum_id=forum_id, classroom_id=classroom_id)
    #fclass.deadline_date = deadline_date.strftime('%d/%m/%Y')
    fclass.deadline_date = datetime.strptime(deadline_date, '%Y %B %d - %I:%M %p')
    fclass.save()
    return JsonResponse({'status':deadline_date}, safe=False)             
        
# 列出文字註記類別
class SpeculationAnnotationListView(ListView):
    model = SpeculationAnnotation
    context_object_name = 'contents'
    template_name = "teacher/speculation_annotation.html"		
    def get_queryset(self):
        queryset = SpeculationAnnotation.objects.filter(forum_id=self.kwargs['forum_id'])
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(SpeculationAnnotationListView, self).get_context_data(**kwargs)
        fwork = SpeculationWork.objects.get(id=self.kwargs['forum_id'])
        context['fwork']= fwork
        context['forum_id'] = self.kwargs['forum_id']
        return context	
			
#新增一個註記類別
class SpeculationAnnotationCreateView(CreateView):
    model = SpeculationAnnotation
    form_class = SpeculationAnnotationForm
    template_name = "teacher/speculation_annotation_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        work = SpeculationAnnotation(forum_id=self.object.forum_id)
        work.kind = self.object.kind
        work.color = self.object.color
        work.save()         
  
        return redirect("/teacher/speculation/annotation/"+self.kwargs['forum_id'])  

    def get_context_data(self, **kwargs):
        ctx = super(SpeculationAnnotationCreateView, self).get_context_data(**kwargs)
        ctx['forum_id'] = self.kwargs['forum_id']
        return ctx

def speculation_annotation_delete(request, forum_id, content_id):
    instance = SpeculationAnnotation.objects.get(id=content_id)
    instance.delete()

    return redirect("/teacher/speculation/annotation/"+forum_id)  
	
def speculation_annotation_edit(request, forum_id, content_id):
    try:
        instance = SpeculationAnnotation.objects.get(id=content_id)
    except:
        pass
    if request.method == 'POST':
            content_id = request.POST.get("content_id", "")
            try:
                content = SpeculationAnnotation.objects.get(id=content_id)
            except ObjectDoesNotExist:
	              content = SpeculattionAnnotation(forum_id= request.POST.get("forum_id", ""))
            content.kind = request.POST.get("kind", "")
            content.color = request.POST.get("color", "")								
            content.save()
            return redirect('/teacher/speculation/annotation/'+forum_id)   
    return render(request,'teacher/speculation_annotation_form.html',{'content': instance, 'forum_id':forum_id, 'content_id':content_id})

def speculation_group(request, classroom_id, forum_id):
    title = SpeculationWork.objects.get(id=forum_id).title
    speculation = SpeculationClass.objects.get(forum_id=forum_id, classroom_id=classroom_id)
    groups = ClassroomGroup.objects.filter(classroom_id=speculation.classroom_id).order_by("-id")
    return render(request,'teacher/speculation_group.html',{'speculation': speculation, 'groups':groups, 'title':title})

def speculation_group_set(request):
    group_id = request.POST.get('groupid')
    forum_id = request.POST.get('forumid')
    classroom_id = request.POST.get('classroomid')
    if group_id and forum_id and classroom_id:      
        forum = SpeculationClass.objects.get(forum_id=forum_id, classroom_id=classroom_id)	
        if is_teacher(request.user, forum.classroom_id) or is_assistant(request.user, forum.classroom_id):
            forum.group = group_id
            forum.save()      
        return JsonResponse({'status':'ok'}, safe=False)
    else:
        return JsonResponse({'status':'fail'}, safe=False) 
# 列出所有教師
class TeacherListView(ListView):
    model = User
    context_object_name = 'teachers'
    template_name = 'teacher/member.html'
    paginate_by = 50
		
    def get_queryset(self):      
        teachers = Group.objects.get(name="teacher").user_set.all().order_by("-last_login")
        queryset = []
        classrooms = Classroom.objects.all()
        fworks = FWork.objects.all()
        sworks = SpeculationWork.objects.all()
        for teacher in teachers:
            rooms = filter(lambda w: w.teacher_id==teacher.id, classrooms)
            classroom_ids = []									
            for classroom in rooms:
                classroom_ids.append(classroom.id)
            enroll = Enroll.objects.filter(classroom_id__in=classroom_ids, seat__gt=0).count()
            fwork = filter(lambda w: w.teacher_id==teacher.id, fworks)
            swork = filter(lambda w: w.teacher_id==teacher.id, sworks)
            queryset.append([teacher, len(rooms), len(fwork), len(swork), enroll])
        return queryset
			
# 列出某教師的所有學生
class StudentListView(ListView):
    model = Classroom
    context_object_name = 'students'
    template_name = 'teacher/student.html'
    paginate_by = 10
		
    def get_queryset(self):      
        queryset = []
        classrooms = Classroom.objects.filter(teacher_id=self.kwargs['teacher_id']).order_by("-id")
        for classroom in classrooms:            
            enrolls = Enroll.objects.filter(classroom_id=classroom.id, seat__gt=0).order_by("seat")
            queryset.append([classroom, enrolls])
        return queryset			
			
# 列出所有組別
class GroupListView(ListView):
    model = Group
    context_object_name = 'groups'
    template_name = 'teacher/group.html'
    paginate_by = 25
    def get_queryset(self):      
        queryset = ClassroomGroup.objects.filter(classroom_id=self.kwargs['classroom_id']).order_by("-id")
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(GroupListView, self).get_context_data(**kwargs)
        context['classroom_id'] = self.kwargs['classroom_id']
        return context				
			
#新增一個分組
class GroupCreateView(CreateView):
    model = ClassroomGroup
    form_class = GroupForm
    template_name = 'teacher/group_form.html'    
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.classroom_id = self.kwargs['classroom_id']
        if is_teacher(self.request.user, self.kwargs['classroom_id']) or is_assistant(self.request.user, self.kwargs['classroom_id']):
            self.object.save()
            # 隨機分組
            if form.cleaned_data['assign'] == 1:
                enrolls = Enroll.objects.filter(classroom_id=self.kwargs['classroom_id'], seat__gt=0).order_by('?')
                number = 0
                for enroll in enrolls:
                    group = StudentGroup(group_id=self.object.id, enroll_id=enroll.id, group=(number % self.object.numbers))
                    group.save()
                    number += 1
                self.object.opening=False
                self.object.save()
                
        return redirect("/student/group/list/"+ str(self.object.id))   
			
    def get_context_data(self, **kwargs):
        context = super(GroupCreateView, self).get_context_data(**kwargs)
        return context	
			
class GroupUpdateView(UpdateView):
    model = ClassroomGroup
    form_class = GroupForm2		
    template_name = 'form.html'
    def get_success_url(self):
        succ_url =  '/student/group/list/'+self.kwargs['pk']
        return succ_url
			
    def form_valid(self, form):
        if is_teacher(self.request.user, self.kwargs['classroom_id']) or is_assistant(self.request.user, self.kwargs['classroom_id']):
            group = ClassroomGroup.objects.get(id=self.kwargs['pk'])
            reduce = group.numbers - form.cleaned_data['numbers']
            if reduce > 0:
                for i in range(reduce):
                    StudentGroup.objects.filter(group_id=self.kwargs['pk'], group=group.numbers-i).delete()
            form.save()
        return HttpResponseRedirect(self.get_success_url())
			

# 分組
def make(request):
    group_id = request.POST.get('groupid')
    action = request.POST.get('action')
    if group_id and action :      
        group = ClassroomGroup.objects.get(id=group_id)	
        if is_teacher(request.user, group.classroom_id) or is_assistant(request.user, group.classroom_id):
            if action == 'open':            
                group.opening = True   
            else : 
                group.opening = False
            group.save()      
        return JsonResponse({'status':'ok'}, safe=False)
    else:
        return JsonResponse({'status':'fail'}, safe=False) 
			
# 分組
def make2(request, group_id, action):
        group = ClassroomGroup.objects.get(id=group_id)	
        if is_teacher(request.user, group.classroom_id) or is_assistant(request.user, group.classroom_id):
            if action == '1':            
                group.opening = True   
            else : 
                group.opening = False
            group.save()      
        return redirect("/student/group/list/"+str(group.id))
			
# 列出所有測驗主題
class ExamListView(ListView):
    model = Exam
    context_object_name = 'exams'
    template_name = "teacher/exam_list.html"		
    paginate_by = 20
    def get_queryset(self):        
        exam_classes = ExamClass.objects.filter(classroom_id=self.kwargs['classroom_id']).order_by("-publication_date", "-exam_id")
        exams = []
        for exam_class in exam_classes:
            exam = Exam.objects.get(id=exam_class.exam_id)
            exams.append([exam, exam_class])
        return exams
			
    def get_context_data(self, **kwargs):
        context = super(ExamListView, self).get_context_data(**kwargs)
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['classroom'] = classroom
        return context	
        
#新增一個測驗主題
class ExamCreateView(CreateView):
    model = Exam
    form_class = ExamForm
    template_name = "teacher/exam_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.user_id = self.request.user.id
        self.object.classroom_id = self.kwargs['classroom_id']
        self.object.domains = self.request.POST.getlist('domains')
        self.object.levels = self.request.POST.getlist('levels')	        
        self.object.save()  
        classrooms = self.request.POST.getlist('classrooms')
        for classroom in classrooms:
          exam_class = ExamClass(exam_id=self.object.id, classroom_id=classroom)
          exam_class.save()
        
        return redirect("/teacher/exam/"+self.kwargs['classroom_id'])           
        
    def get_context_data(self, **kwargs):
        context = super(ExamCreateView, self).get_context_data(**kwargs)
        classroom_list = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id)
        for classroom in classrooms:
            classroom_list.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            if not assistant.classroom_id in classroom_list:
                classroom_list.append(assistant.classroom_id)
        classrooms = Classroom.objects.filter(id__in=classroom_list).order_by("-id")
        context['classrooms'] = classrooms
        context['classroom_id'] = int(self.kwargs['classroom_id'])
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['domains'] = Domain.objects.all()
        context['levels'] = Level.objects.all()
        return context	
  
        return redirect("/teacher/exam/"+self.kwargs['classroom_id']) 			
	
	
def exam_categroy(request, classroom_id, exam_id):
    exam = Exam.objects.get(id=exam_id)
    domains = Domain.objects.all()
    levels = Level.objects.all()		
    if request.method == 'POST':
        form = ExamCategroyForm(request.POST)
        if form.is_valid():
            exam.domains = request.POST.getlist('domains')
            exam.levels = request.POST.getlist('levels')	
            exam.save()
            return redirect('/teacher/exam/'+classroom_id+'/#'+str(exam.id))
    else:
        form = ExamCategroyForm(instance=exam)
        
    return render(request,'teacher/exam_categroy_form.html',{'domains': domains, 'levels':levels, 'classroom_id': classroom_id, 'exam':exam})

	
# 列出所有討論主題
class ExamAllListView(ListView):
    model = Exam
    context_object_name = 'exams'
    template_name = "teacher/exam_all.html"		
    paginate_by = 20
		
    def get_queryset(self):
      # 年級
      if self.kwargs['categroy'] == "1":
        queryset = FWork.objects.filter(levels__contains=self.kwargs['categroy_id']).order_by("-id")
      # 學習領域
      elif self.kwargs['categroy'] == "2":
        queryset = FWork.objects.filter(domains__contains=self.kwargs['categroy_id']).order_by("-id")   
      else:
        queryset = FWork.objects.all().order_by("-id")
      if self.request.GET.get('account') != None:
        keyword = self.request.GET.get('account')
        users = User.objects.filter(Q(username__icontains=keyword) | Q(first_name__icontains=keyword)).order_by('-id')
        user_list = []
        for user in users:
            user_list.append(user.id)
        forums = queryset.filter(teacher_id__in=user_list)
        return forums
      else:				
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(ExamAllListView, self).get_context_data(**kwargs)
        context['categroy'] = self.kwargs['categroy']							
        context['categroy_id'] = self.kwargs['categroy_id']							
        context['levels'] = Level.objects.all()				
        context['domains'] = Domain.objects.all()
        return context	

# 列出某測驗主題的班級
class ExamClassListView(ListView):
    model = Exam
    context_object_name = 'classrooms'
    template_name = "teacher/exam_class.html"		
    paginate_by = 20
	
    def get_queryset(self):        		
        eclass_dict = dict(((eclass.classroom_id, eclass) for eclass in ExamClass.objects.filter(exam_id=self.kwargs['exam_id'])))		
        classroom_list = []
        classroom_ids = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id).order_by("-id")
        for classroom in classrooms:
            if classroom.id in eclass_dict:
                classroom_list.append([classroom, True, eclass_dict[classroom.id].deadline, eclass_dict[classroom.id].deadline_date])
            else :
                classroom_list.append([classroom, False, False, timezone.now()])
            classroom_ids.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            classroom = Classroom.objects.get(id=assistant.classroom_id)
            if not classroom.id in classroom_ids:
                if classroom.id in eclass_dict:
                    classroom_list.append([classroom, True, eclass_dict[classroom.id].deadline, eclass_dict[classroom.id].deadline_date])
                else :
                    classroom_list.append([classroom, False, False, timezone.now()])
        return classroom_list
			
    def get_context_data(self, **kwargs):
        context = super(ExamClassListView, self).get_context_data(**kwargs)				
        exam = Exam.objects.get(id=self.kwargs['exam_id'])
        context['exam'] = exam
        context['exam_id'] = self.kwargs['exam_id']
        return context	
	
# Ajax 開放班取、關閉班級
def exam_switch(request):
    exam_id = request.POST.get('examid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        examclass = ExamClass.objects.get(exam_id=exam_id, classroom_id=classroom_id)
        if status == 'false' :
    				examclass.delete()
    except ObjectDoesNotExist:
        if status == 'true':
            examclass = ExamClass(exam_id=exam_id, classroom_id=classroom_id)
            examclass.save()
    return JsonResponse({'status':status}, safe=False)        
	
class ExamEditUpdateView(UpdateView):
    model = Exam
    fields = ['title']
    template_name = 'form.html'
    #success_url = '/teacher/forum/domain/'
    def get_success_url(self):
        succ_url =  '/teacher/exam/'+self.kwargs['classroom_id']
        return succ_url
		
def exam_deadline(request, classroom_id, exam_id):
    exam = Exam.objects.get(id=exam_id)
    if request.method == 'POST':
        form = ExamCategroyForm(request.POST)
        if form.is_valid():
            exam.domains = request.POST.getlist('domains')
            exam.levels = request.POST.getlist('levels')	
            forum.save()
            return redirect('/teacher/exam/'+classroom_id)
    else:
        examclass = ExamClass.objects.get(classroom_id=classroom_id, exam_id=exam_id)
        form = ExamDeadlineForm(instance=examclass)
    return render(request,'teacher/exam_deadline_form.html',{'examclass':examclass})

	
# Ajax 設定期限、取消期限
def exam_deadline_set(request):
    exam_id = request.POST.get('examid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        examclass = ExamClass.objects.get(exam_id=exam_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        examclass = Examclass(exam_id=exam_id, classroom_id=classroom_id)
    if status == 'True':
        examclass.deadline = True
    else :
        examclass.deadline = False
    examclass.save()
    return JsonResponse({'status':status}, safe=False)        

# Ajax 設定期限日期
def exam_deadline_date(request):
    exam_id = request.POST.get('examid')
    classroom_id = request.POST.get('classroomid')		
    deadline_date = request.POST.get('deadlinedate')
    try:
        examclass = ExamClass.objects.get(exam_id=exam_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        examclass = ExamClass(exam_id=exam_id, classroom_id=classroom_id)
    #fclass.deadline_date = deadline_date.strftime('%d/%m/%Y')
    examclass.deadline_date = datetime.strptime(deadline_date, '%Y %B %d - %I:%M %p')
    examclass.save()
    return JsonResponse({'status':deadline_date}, safe=False)             
		
# 列出所有測驗題目
class ExamQuestionListView(ListView):
    model = ExamQuestion
    context_object_name = 'questions'
    template_name = "teacher/exam_question.html"		
    def get_queryset(self):
        queryset = ExamQuestion.objects.filter(exam_id=self.kwargs['exam_id']).order_by("-id")
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(ExamQuestionListView, self).get_context_data(**kwargs)
        exam = Exam.objects.get(id=self.kwargs['exam_id'])
        context['exam']= exam
        context['exam_id'] = self.kwargs['exam_id']
        questions = ExamQuestion.objects.filter(exam_id=self.kwargs['exam_id'])
        context['score_total'] = sum(question.score for question in questions)			
        return context	
			
#新增一個題目
class ExamQuestionCreateView(CreateView):
    model = ExamQuestion
    form_class = ExamQuestionForm
    template_name = "teacher/exam_question_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        question = ExamQuestion(exam_id=self.object.exam_id)
        question.answer = self.object.answer
           
        if 'title_pic' in self.request.FILES :
            myfile = self.request.FILES['title_pic']
            fs = FileSystemStorage()
            filename = uuid4().hex
            question.title_filename = str(self.request.user.id)+"/"+filename
            fs.save("static/exam/"+str(self.request.user.id)+"/"+filename, myfile)                
        question.title = self.object.title
        question.types = self.object.types
        question.score = self.object.score
        question.save()         
  
        return redirect("/teacher/exam/question/"+self.kwargs['exam_id'])  

    def get_context_data(self, **kwargs):
        ctx = super(ExamQuestionCreateView, self).get_context_data(**kwargs)
        ctx['exam'] = Exam.objects.get(id=self.kwargs['exam_id'])
        return ctx

def exam_question_delete(request, exam_id, question_id):
    instance = ExamQuestion.objects.get(id=question_id)
    instance.delete()

    return redirect("/teacher/exam/question/"+exam_id)  
	
def exam_question_edit(request, exam_id, question_id):
    exam = Exam.objects.get(id=exam_id)
    try:
        instance = ExamQuestion.objects.get(id=question_id)
    except:
        pass
    if request.method == 'POST':
            question_id = request.POST.get("question_id", "")
            try:
                question = ExamQuestion.objects.get(id=question_id)
            except ObjectDoesNotExist:
	              question = ExamQuestion(exam_id= request.POST.get("exam_id", ""), types=form.cleaned_data['types'])
            if question.types == 2:
                question.option1 = request.POST.get("option1", "")	
                question.option2 = request.POST.get("option2", "")	
                question.option3 = request.POST.get("option3", "")	
                question.option4 = request.POST.get("option4", "")	
            question.score = request.POST.get("score", "")
            question.answer = request.POST.get("answer", "")	
            question.title = request.POST.get("title", "")
            if 'title_pic' in request.FILES :
                myfile = request.FILES['title_pic']
                fs = FileSystemStorage()
                filename = uuid4().hex
                question.title_filename = str(request.user.id)+"/"+filename
                fs.save("static/exam/"+str(request.user.id)+"/"+filename, myfile)  
            question.save()
            return redirect('/teacher/exam/question/'+exam_id+"#"+str(question.id))   
    return render(request,'teacher/exam_question_edit.html',{'question': instance, 'exam':exam, 'quesiton_id':question_id})	

# Create your views here.
def exam_import_sheet(request, types, exam_id):
    #if request.user.id != 1:
    #    return redirect("/")
    if request.method == "POST":
        form = UploadFileForm(request.POST,
                              request.FILES)
        if form.is_valid():
            ExamImportQuestion.objects.all().delete()            
            if types == "1":
                request.FILES['file'].save_to_database(
                    name_columns_by_row=0,
                    model=ExamImportQuestion,
                    mapdict=['title', 'answer', 'score'])
            elif types == "2":
                request.FILES['file'].save_to_database(
                    name_columns_by_row=0,
                    model=ExamImportQuestion,
                    mapdict=['title', 'option1', 'option2','option3','option4','answer', 'score'])                    
            elif types == "3":
                request.FILES['file'].save_to_database(
                    name_columns_by_row=0,
                    model=ExamImportQuestion,
                    mapdict=['title', 'answer', 'score'])
            questions = ExamImportQuestion.objects.all()
            return render(request, 'teacher/exam_import_question.html',{'types':types, 'questions':questions, 'exam_id': exam_id})
        else:
            return HttpResponseBadRequest()
    else:	
        form = UploadFileForm()
    return render(
        request,
        'teacher/exam_upload_form.html',
        {
					  'exam_id': exam_id, 
            'form': form,
            'title': 'Excel file upload and download example',
            'header': ('Please choose any excel file ' +
                       'from your cloned repository:')
        })
	
# Create your views here.
def exam_import_question(request, types, exam_id):
    #if request.user.id != 1:
    #    return redirect("/")
    questions = ExamImportQuestion.objects.all()           
    if types == "1":
        for question in questions:            
            new_question = ExamQuestion(exam_id=exam_id, types=1, title=question.title, answer=question.answer, score=question.score)        
            new_question.save()
    elif types == "2":
        for question in questions:         
            new_question = ExamQuestion(exam_id=exam_id, types=2, title=question.title, option1=question.option1, option2=question.option2, option3=question.option3, option4=question.option4, answer=question.answer, score=question.score)
            new_question.save()
    elif types == "3":
        for question in questions:     
            new_question = ExamQuestion(exam_id=exam_id, types=3, title=question.title, answer=question.answer, score=question.score)
            new_question.save()
            
    return redirect('/teacher/exam/question/'+exam_id)			
	
def exam_round(request, classroom_id, exam_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    examclass = ExamClass.objects.get(classroom_id=classroom_id, exam_id=exam_id)
    return render(request,'teacher/exam_round.html',{'examclass':examclass})		
	
def exam_round_set(request):
    exam_id = request.POST.get('examid')
    classroom_id = request.POST.get('classroomid')		
    round_limit = request.POST.get('round_limit')
    try:
        examclass = ExamClass.objects.get(exam_id=exam_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        examclass = Examclass(exam_id=exam_id, classroom_id=classroom_id)
    examclass.round_limit = int(round_limit)
    examclass.save()
    return JsonResponse({'status':'ok'}, safe=False)  	
	
def exam_score(request, classroom_id, exam_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    exam = Exam.objects.get(id=exam_id)
    classroom = Classroom.objects.get(id=classroom_id)
    examclass = ExamClass.objects.get(classroom_id=classroom_id, exam_id=exam_id)
    enrolls = Enroll.objects.filter(classroom_id=classroom_id, seat__gt=0).order_by("seat")
    enroll_ids = []
    for enroll in enrolls:
        enroll_ids.append(enroll.student_id)
    examworks = ExamWork.objects.filter(exam_id=exam_id, student_id__in=enroll_ids, publish=True).order_by("-id")
    scores = []
    for enroll in enrolls:
        works = filter(lambda w: w.student_id == enroll.student_id, examworks)
        if len(works) > 0 :
            score_max = max(work.score for work in works)
            score_avg = sum(work.score for work in works) / len(works)	
        else :
            score_max = 0
            score_avg = 0
        scores.append([enroll, works, score_avg, score_max])
    return render(request,'teacher/exam_score.html',{'classroom': classroom, 'exam':exam, 'scores':scores})		
	
# 列出所有討論測驗
class ExamAllListView(ListView):
    model = Exam
    context_object_name = 'exams'
    template_name = "teacher/exam_all.html"		
    paginate_by = 20
		
    def get_queryset(self):
      # 年級
      if self.kwargs['categroy'] == "1":
        queryset = Exam.objects.filter(levels__contains=self.kwargs['categroy_id']).order_by("-id")
      # 學習領域
      elif self.kwargs['categroy'] == "2":
        queryset = Exam.objects.filter(domains__contains=self.kwargs['categroy_id']).order_by("-id")   
      else:
        queryset = Exam .objects.all().order_by("-id")
      if self.request.GET.get('account') != None:
        keyword = self.request.GET.get('account')
        users = User.objects.filter(Q(username__icontains=keyword) | Q(first_name__icontains=keyword)).order_by('-id')
        user_list = []
        for user in users:
            user_list.append(user.id)
        exams = queryset.filter(teacher_id__in=user_list)
        return exams
      else:				
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(ExamAllListView, self).get_context_data(**kwargs)
        context['categroy'] = self.kwargs['categroy']							
        context['categroy_id'] = self.kwargs['categroy_id']							
        context['levels'] = Level.objects.all()				
        context['domains'] = Domain.objects.all()
        return context	

       

# 列出所有討論主題
class TeamListView(ListView):
    model = TeamWork
    context_object_name = 'teams'
    template_name = "teacher/team_list.html"		
    paginate_by = 20
    def get_queryset(self):        
        teamclasses = TeamClass.objects.filter(classroom_id=self.kwargs['classroom_id']).order_by("-publication_date", "-team_id")
        teams = []
        for teamclass in teamclasses:
            team = TeamWork.objects.get(id=teamclass.team_id)
            teams.append([team, teamclass])
        return teams
			
    def get_context_data(self, **kwargs):
        context = super(TeamListView, self).get_context_data(**kwargs)
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['classroom'] = classroom
        return context	
        
#新增一個討論主題
class TeamCreateView(CreateView):
    model = TeamWork
    form_class = TeamForm
    template_name = "teacher/team_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.teacher_id = self.request.user.id
        self.object.classroom_id = self.kwargs['classroom_id']
        self.object.domains = self.request.POST.getlist('domains')
        self.object.levels = self.request.POST.getlist('levels')	        
        self.object.save()  
        classrooms = self.request.POST.getlist('classrooms')
        for classroom in classrooms:
          teamclass = TeamClass(team_id=self.object.id, classroom_id=classroom)
          teamclass.save()
        
        return redirect("/teacher/team/"+self.kwargs['classroom_id'])           
        
    def get_context_data(self, **kwargs):
        context = super(TeamCreateView, self).get_context_data(**kwargs)
        classroom_list = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id)
        for classroom in classrooms:
            classroom_list.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            if not assistant.classroom_id in classroom_list:
                classroom_list.append(assistant.classroom_id)
        classrooms = Classroom.objects.filter(id__in=classroom_list).order_by("-id")
        context['classrooms'] = classrooms
        context['classroom_id'] = int(self.kwargs['classroom_id'])
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['domains'] = Domain.objects.all()
        context['levels'] = Level.objects.all()
        return context	
	
			
def team_categroy(request, classroom_id, team_id):
    team = TeamWork.objects.get(id=team_id)
    domains = Domain.objects.all()
    levels = Level.objects.all()		
    if request.method == 'POST':
        form = TeamCategroyForm(request.POST)
        if form.is_valid():
            team.domains = request.POST.getlist('domains')
            team.levels = request.POST.getlist('levels')	
            team.save()
            return redirect('/teacher/team/'+classroom_id+'/#'+str(team.id))
    else:
        form = TeamCategroyForm(instance=team)
        
    return render(request,'teacher/team_categroy_form.html',{'domains': domains, 'levels':levels, 'classroom_id': classroom_id, 'team':team})


# 列出某任務主題的班級
class TeamClassListView(ListView):
    model = TeamWork
    context_object_name = 'classrooms'
    template_name = "teacher/team_class.html"		
    paginate_by = 20
	
    def get_queryset(self):        		
        teamclass_dict = dict(((teamclass.classroom_id, teamclass) for teamclass in TeamClass.objects.filter(team_id=self.kwargs['team_id'])))		
        classroom_list = []
        classroom_ids = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id).order_by("-id")
        for classroom in classrooms:
            if classroom.id in teamclass_dict:
                classroom_list.append([classroom, True, teamclass_dict[classroom.id].deadline, teamclass_dict[classroom.id].deadline_date])
            else :
                classroom_list.append([classroom, False, False, timezone.now()])
            classroom_ids.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            classroom = Classroom.objects.get(id=assistant.classroom_id)
            if not classroom.id in classroom_ids:
                if classroom.id in teamclass_dict:
                    classroom_list.append([classroom, True, teamclass_dict[classroom.id].deadline, teamclass_dict[classroom.id].deadline_date])
                else :
                    classroom_list.append([classroom, False, False, timezone.now()])
        return classroom_list
			
    def get_context_data(self, **kwargs):
        context = super(TeamClassListView, self).get_context_data(**kwargs)				
        teamwork = TeamWork.objects.get(id=self.kwargs['team_id'])
        context['teamwork'] = teamwork
        context['team_id'] = self.kwargs['team_id']
        return context	
	
# Ajax 開放班取、關閉班級
def team_switch(request):
    team_id = request.POST.get('teamid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        teamwork = TeamClass.objects.get(team_id=team_id, classroom_id=classroom_id)
        if status == 'false' :
    				teamwork.delete()
    except ObjectDoesNotExist:
        if status == 'true':
            teamwork = TeamClass(team_id=team_id, classroom_id=classroom_id)
            teamwork.save()
    return JsonResponse({'status':status}, safe=False)        
	
# 列出某任務所有同學名單
def team_class(request, classroom_id, work_id):
    enrolls = Enroll.objects.filter(classroom_id=classroom_id)
    classroom_name = Classroom.objects.get(id=classroom_id).name
    classmate_work = []
    scorer_name = ""
    for enroll in enrolls:
        try:    
            work = SWork.objects.get(student_id=enroll.student_id, index=work_id)
            if work.scorer > 0 :
                scorer = User.objects.get(id=work.scorer)
                scorer_name = scorer.first_name
            else :
                scorer_name = "1"
        except ObjectDoesNotExist:
            work = SWork(index=work_id, student_id=1)
        try:
            group_name = EnrollGroup.objects.get(id=enroll.group).name
        except ObjectDoesNotExist:
            group_name = "沒有組別"
        assistant = Assistant.objects.filter(classroom_id=classroom_id, student_id=enroll.student_id, lesson=work_id)
        if assistant.exists():
            classmate_work.append([enroll,work,1, scorer_name, group_name])
        else :
            classmate_work.append([enroll,work,0, scorer_name, group_name])   
    def getKey(custom):
        return custom[0].seat
	
    classmate_work = sorted(classmate_work, key=getKey)
   
    return render(request,'teacher/twork_class.html',{'classmate_work': classmate_work, 'classroom_id':classroom_id, 'index': work_id})
			
def team_deadline(request, classroom_id, team_id):
    team = TeamWork.objects.get(id=team_id)
    classroom = Classroom.objects.get(id=classroom_id)
    if request.method == 'POST':
        form = CategroyForm(request.POST)
        if form.is_valid():
            team.domains = request.POST.getlist('domains')
            team.levels = request.POST.getlist('levels')	
            team.save()
            return redirect('/teacher/team/'+classroom_id)
    else:
        teamclass = TeamClass.objects.get(classroom_id=classroom_id, team_id=team_id)
        form = TeamDeadlineForm(instance=teamclass)
        teamclasses = TeamClass.objects.filter(team_id=team_id).order_by("-id")
    return render(request,'teacher/team_deadline_form.html',{'teamclasses':teamclasses, 'teamclass':teamclass, 'team':team, 'classroom':classroom})

# Ajax 設定期限、取消期限
def team_deadline_set(request):
    team_id = request.POST.get('teamid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        teamclass = TeamClass.objects.get(team_id=team_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        teamclass = Teamclass(team_id=team_id, classroom_id=classroom_id)
    if status == 'True':
        teamclass.deadline = True
    else :
        teamclass.deadline = False
    teamclass.save()
    return JsonResponse({'status':status}, safe=False)        

# Ajax 設定期限日期
def team_deadline_date(request):
    team_id = request.POST.get('teamid')
    classroom_id = request.POST.get('classroomid')		
    deadline_date = request.POST.get('deadlinedate')
    try:
        teamclass = TeamClass.objects.get(team_id=team_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        teamclass = TeamClass(team_id=team_id, classroom_id=classroom_id)
    #fclass.deadline_date = deadline_date.strftime('%d/%m/%Y')
    teamclass.deadline_date = datetime.strptime(deadline_date, '%Y %B %d - %H:%M')
    teamclass.save()
    return JsonResponse({'status':deadline_date}, safe=False)            		
	
			
class TeamEditUpdateView(UpdateView):
    model = TeamWork
    fields = ['title']
    template_name = 'form.html'
    #success_url = '/teacher/forum/domain/'
    def get_success_url(self):
        succ_url =  '/teacher/team/'+self.kwargs['classroom_id']
        return succ_url
			
def team_group(request, classroom_id, team_id):
    enrolls = Enroll.objects.filter(classroom_id=classroom_id)
    enroll_dict = {}
    for enroll in enrolls:
        enroll_dict[enroll.id] = enroll
    groups = ClassroomGroup.objects.filter(classroom_id=classroom_id)
    group_ids = []
    for group in groups:
        group_ids.append(group.id)
    classroom = Classroom.objects.get(id=classroom_id)
    studentgroups = StudentGroup.objects.filter(group_id__in=group_ids)
    group_list = []
    for group in groups:        
        groupclass_list = []  
        groupclass_dict = {}
        students = filter(lambda student: student.group_id == group.id , studentgroups)
        for student in students:
            if student.enroll_id in enroll_dict:
                if student.group in groupclass_dict:
                    groupclass_dict[student.group].append(enroll_dict[student.enroll_id])
                else :
                    groupclass_dict[student.group] = [enroll_dict[student.enroll_id]]
        for key in groupclass_dict:
            groupclass_list.append([key, groupclass_dict[key]])
        group_list.append([group.id, groupclass_list])
    teamclass = TeamClass.objects.get(team_id=team_id, classroom_id=classroom_id)
    try:
        group = ClassroomGroup.objects.get(id=teamclass.group)
    except ObjectDoesNotExist:
        group = ClassroomGroup(title="不分組", id=0)
    return render(request,'teacher/team_group.html',{'team_id': team_id, 'teamgroup': group, 'groups':groups, 'classroom':classroom, 'group_list':group_list})

# 影片觀看時間統計
class EventVideoView(ListView):
    context_object_name = 'events'
    #paginate_by = 50
    template_name = 'teacher/event_video.html'

    def get_queryset(self):    
				enrolls = Enroll.objects.filter(classroom_id=self.kwargs['classroom_id'], seat__gt=0).order_by("seat")
				events = []
				for enroll in enrolls: 
						videos = VideoLogHelper().getLogByUserid(enroll.student_id,self.kwargs['work_id'])
						length = 0
						for video in videos: 
										length += video['length']
						events.append([enroll, length/60.0])
				return events
			
    def get_context_data(self, **kwargs):
        context = super(EventVideoView, self).get_context_data(**kwargs)
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['length'] = FContent.objects.get(id=self.kwargs['work_id']).youtube_length / 60.0
        context['content_id'] = self.kwargs['work_id']
        context['classroom'] = classroom
        enrolls = Enroll.objects.filter(classroom_id=classroom.id)
        context['height'] = 100 + enrolls.count() * 40
        return context
			
# 記錄影片長度
def video_length(request):
    content_id = request.POST.get('content_id')
    length = request.POST.get('length')
    page = request.POST.get('page')
    if page == "team":
        teamcontent = TeamContent.objects.get(id=content_id)
        teamcontent.youtube_length = length
        teamcontent.save()
    elif page == "forum":
        fcontent = FContent.objects.get(id=content_id)
        fcontent.youtube_length = length
        fcontent.save()
    elif page == "couurse":
        coursecontent = CourseContent.objects.get(id=content_id)
        coursecontent.youtube_length = length
        coursecontent.save()        
    return JsonResponse({'status':'ok'}, safe=False)	
	
# 影片記錄條
class VideoListView(ListView):
    context_object_name = 'videos'
    template_name = 'teacher/event_video_user.html'
    
    def get_queryset(self):
				videos = VideoLogHelper().getLogByUserid(self.kwargs['user_id'],self.kwargs['content_id'])        
				return videos
        
    def get_context_data(self, **kwargs):
        context = super(VideoListView, self).get_context_data(**kwargs)
        content = FContent.objects.get(id=self.kwargs['content_id'])
        context['user_id'] = self.kwargs['user_id']
        context['content'] = content
        context['length'] = content.youtube_length
        return context  

    # 限本班任課教師或助教     
    def render_to_response(request,self, context):
        if not is_teacher(self.request.user ,self.kwargs['classroom_id']):
            if not is_assistant(self.request.user, self.kwargs['classroom_id'] ):
                  return redirect('/')
        return super(VideoListView, self).render(request,context)   
			
# Ajax 設定合作區組別
def team_group_set(request):
    team_id = request.POST.get('teamid')
    classroom_id = request.POST.get('classroomid')		
    group = request.POST.get('groupid')
    try:
        teamclass = TeamClass.objects.get(team_id=int(team_id), classroom_id=int(classroom_id))
    except ObjectDoesNotExist:
        pass
    teamclass.group = int(group)
    teamclass.save()    
    return JsonResponse({'status':team_id}, safe=False)  

# 列出所有討論主題
class CourseListView(ListView):
    model = CourseWork
    context_object_name = 'courses'
    template_name = "teacher/course_list.html"		
    paginate_by = 20
    def get_queryset(self):        
        courseclasses = CourseClass.objects.filter(classroom_id=self.kwargs['classroom_id']).order_by("-publication_date", "-course_id")
        courses = []
        for courseclass in courseclasses:
            course = CourseWork.objects.get(id=courseclass.course_id)
            courses.append([course, courseclass])
        return courses
			
    def get_context_data(self, **kwargs):
        context = super(CourseListView, self).get_context_data(**kwargs)
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['classroom'] = classroom
        return context	
        
#新增一個討論主題
class CourseCreateView(CreateView):
    model = CourseWork
    form_class = CourseForm
    template_name = "teacher/course_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.teacher_id = self.request.user.id
        self.object.classroom_id = self.kwargs['classroom_id']
        self.object.domains = self.request.POST.getlist('domains')
        self.object.levels = self.request.POST.getlist('levels')	        
        self.object.save()  
        classrooms = self.request.POST.getlist('classrooms')
        for classroom in classrooms:
          course_class = CourseClass(course_id=self.object.id, classroom_id=classroom)
          course_class.save()
        
        return redirect("/teacher/course/"+self.kwargs['classroom_id'])           
        
    def get_context_data(self, **kwargs):
        context = super(CourseCreateView, self).get_context_data(**kwargs)
        classroom_list = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id)
        for classroom in classrooms:
            classroom_list.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            if not assistant.classroom_id in classroom_list:
                classroom_list.append(assistant.classroom_id)
        classrooms = Classroom.objects.filter(id__in=classroom_list).order_by("-id")
        context['classrooms'] = classrooms
        context['classroom_id'] = int(self.kwargs['classroom_id'])
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['domains'] = Domain.objects.all()
        context['levels'] = Level.objects.all()
        return context	
  
        return redirect("/teacher/course/"+self.kwargs['classroom_id'])        
	
def course_categroy(request, classroom_id, course_id):
    course = CourseWork.objects.get(id=course_id)
    domains = Domain.objects.all()
    levels = Level.objects.all()		
    if request.method == 'POST':
        form = CourseCategroyForm(request.POST)
        if form.is_valid():
            course.domains = request.POST.getlist('domains')
            course.levels = request.POST.getlist('levels')	
            course.save()
            return redirect('/teacher/course/'+classroom_id+'/#'+str(course.id))
    else:
        form = CourseCategroyForm(instance=course)
        
    return render(request,'teacher/course_categroy_form.html',{'domains': domains, 'levels':levels, 'classroom_id': classroom_id, 'course':course})

	
# 列出所有討論主題
class CourseAllListView(ListView):
    model = CourseWork
    context_object_name = 'courses'
    template_name = "teacher/course_all.html"		
    paginate_by = 20
		
    def get_queryset(self):
      # 年級
      if self.kwargs['categroy'] == "1":
        queryset = FWork.objects.filter(levels__contains=self.kwargs['categroy_id']).order_by("-id")
      # 學習領域
      elif self.kwargs['categroy'] == "2":
        queryset = FWork.objects.filter(domains__contains=self.kwargs['categroy_id']).order_by("-id")   
      else:
        queryset = FWork.objects.all().order_by("-id")
      if self.request.GET.get('account') != None:
        keyword = self.request.GET.get('account')
        users = User.objects.filter(Q(username__icontains=keyword) | Q(first_name__icontains=keyword)).order_by('-id')
        user_list = []
        for user in users:
            user_list.append(user.id)
        courses = queryset.filter(teacher_id__in=user_list)
        return courses
      else:				
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(CourseAllListView, self).get_context_data(**kwargs)
        context['categroy'] = self.kwargs['categroy']							
        context['categroy_id'] = self.kwargs['categroy_id']							
        context['levels'] = Level.objects.all()				
        context['domains'] = Domain.objects.all()
        return context	

# 展示討論素材
def course_show(request, course_id):
    course = CourseWork.objects.get(id=course_id)
    domains = Domain.objects.all()
    domain_dict = {}
    for domain in domains :
        key = domain.id
        domain_dict[key] = domain
    levels = Level.objects.all()	
    level_dict = {}
    for level in levels :
        key = level.id
        level_dict[key] = level
    contents = CourseContent.objects.filter(course_id=course_id)
    domains = []		
    if course.domains:
        course_domains = ast.literal_eval(course.domains)
        for domain in course_domains:
            key = int(domain)
            domains.append(domain_dict[key])
    levels = []						
    if course.levels:
        course_levels = ast.literal_eval(course.levels)
        for level in course_levels:
            key = int(level)			
            levels.append(level_dict[key])
    return render(request,'teacher/course_show.html',{'domains':domains, 'levels':levels, 'contents':contents, 'course':course})

		
# 列出某討論主題的班級
class CourseClassListView(ListView):
    model = CourseWork
    context_object_name = 'classrooms'
    template_name = "teacher/course_class.html"		
    paginate_by = 20
	
    def get_queryset(self):        		
        courseclass_dict = dict(((courseclass.classroom_id, courseclass) for courseclass in CourseClass.objects.filter(course_id=self.kwargs['course_id'])))		
        classroom_list = []
        classroom_ids = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id).order_by("-id")
        for classroom in classrooms:
            if classroom.id in courseclass_dict:
                classroom_list.append([classroom, True, courseclass_dict[classroom.id].deadline, courseclass_dict[classroom.id].deadline_date])
            else :
                classroom_list.append([classroom, False, False, timezone.now()])
            classroom_ids.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            classroom = Classroom.objects.get(id=assistant.classroom_id)
            if not classroom.id in classroom_ids:
                if classroom.id in courseclass_dict:
                    classroom_list.append([classroom, True, courseclass_dict[classroom.id].deadline, courseclass_dict[classroom.id].deadline_date])
                else :
                    classroom_list.append([classroom, False, False, timezone.now()])
        return classroom_list
			
    def get_context_data(self, **kwargs):
        context = super(CourseClassListView, self).get_context_data(**kwargs)				
        coursework = CourseWork.objects.get(id=self.kwargs['course_id'])
        context['coursework'] = coursework
        context['course_id'] = self.kwargs['course_id']
        return context	
	
# Ajax 開放班取、關閉班級
def course_switch(request):
    course_id = request.POST.get('courseid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        coursework = CourseClass.objects.get(course_id=course_id, classroom_id=classroom_id)
        if status == 'false' :
    		coursework.delete()
    except ObjectDoesNotExist:
        if status == 'true':
            coursework = CourseClass(course_id=course_id, classroom_id=classroom_id)
            coursework.save()
    return JsonResponse({'status':status}, safe=False)        
	
# 列出某作業所有同學名單
def course_class(request, classroom_id, work_id):
    enrolls = Enroll.objects.filter(classroom_id=classroom_id)
    classroom_name = Classroom.objects.get(id=classroom_id).name
    classmate_work = []
    scorer_name = ""
    for enroll in enrolls:
        try:    
            work = SWork.objects.get(student_id=enroll.student_id, index=work_id)
            if work.scorer > 0 :
                scorer = User.objects.get(id=work.scorer)
                scorer_name = scorer.first_name
            else :
                scorer_name = "1"
        except ObjectDoesNotExist:
            work = SWork(index=work_id, student_id=1)
        try:
            group_name = EnrollGroup.objects.get(id=enroll.group).name
        except ObjectDoesNotExist:
            group_name = "沒有組別"
        assistant = Assistant.objects.filter(classroom_id=classroom_id, student_id=enroll.student_id, lesson=work_id)
        if assistant.exists():
            classmate_work.append([enroll,work,1, scorer_name, group_name])
        else :
            classmate_work.append([enroll,work,0, scorer_name, group_name])   
    def getKey(custom):
        return custom[0].seat
	
    classmate_work = sorted(classmate_work, key=getKey)
   
    return render(request,'teacher/twork_class.html',{'classmate_work': classmate_work, 'classroom_id':classroom_id, 'index': work_id})

# 列出所有討論主題素材
class CourseContentListView(ListView):
    model = CourseContent
    context_object_name = 'contents'
    template_name = "teacher/course_content.html"		
    def get_queryset(self):
        contents = CourseContent.objects.filter(course_id=self.kwargs['course_id']).order_by("id")
        queryset = []
        for content in contents :
            exercises = CourseExercise.objects.filter(content_id=content.id)
            queryset.append([content, exercises])
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(CourseContentListView, self).get_context_data(**kwargs)
        coursework = CourseWork.objects.get(id=self.kwargs['course_id'])
        courseclasses = CourseClass.objects.filter(course_id=self.kwargs['course_id'])				
        context['coursework']= coursework
        context['course_id'] = self.kwargs['course_id']
        context['courseclasses'] = courseclasses
        context['classroom_id'] = self.kwargs['classroom_id']
        return context	
			
#新增一個課程
class CourseContentCreateView(CreateView):
    model = CourseContent
    form_class = CourseContentForm
    template_name = "teacher/course_content_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        work = CourseContent(course_id=self.object.course_id)
        if self.object.types == 1:
            work.types = 1
            work.title = self.object.title
            work.link = self.object.link
        if self.object.types  == 2:
            work.types = 2					
            work.youtube = self.object.youtube
        if self.object.types  == 3:
            work.types = 3
            myfile = self.request.FILES['content_file']
            fs = FileSystemStorage()
            filename = uuid4().hex
            work.title = myfile.name
            work.filename = str(self.request.user.id)+"/"+filename
            fs.save("static/upload/"+str(self.request.user.id)+"/"+filename, myfile)
        if self.object.types  == 4:
            work.types = 4
        work.memo = self.object.memo
        work.save()         
  
        return redirect("/teacher/course/content/"+self.kwargs['classroom_id']+"/"+self.kwargs['course_id'])  

    def get_context_data(self, **kwargs):
        ctx = super(CourseContentCreateView, self).get_context_data(**kwargs)
        ctx['course'] = CourseWork.objects.get(id=self.kwargs['course_id'])
        ctx['classroom_id'] = self.kwargs['classroom_id']        
        return ctx

def course_delete(request, classroom_id, course_id, content_id):
    instance = CourseContent.objects.get(id=content_id)
    instance.delete()

    return redirect("/teacher/course/content/"+classroom_id+"/"+course_id)  
	
def course_edit(request, classroom_id, course_id, content_id):
    try:
        instance = CourseContent.objects.get(id=content_id)
    except:
        pass
    if request.method == 'POST':
            content_id = request.POST.get("id", "")
            try:
                content = CourseContent.objects.get(id=content_id)
            except ObjectDoesNotExist:
	            content = CourseContent(course_id= 0, types=form.cleaned_data['types'])
            if content.types == 1:
                content.title = request.POST.get("title", "")
                content.link = request.POST.get("link", "")
            elif content.types == 2:
                content.youtube = request.POST.get("youtube", "")
            elif content.types == 3:
                myfile =  request.FILES.get("content_file", "")
                fs = FileSystemStorage()
                filename = uuid4().hex
                content.title = myfile.name
                content.filename = str(request.user.id)+"/"+filename
                fs.save("static/upload/"+str(request.user.id)+"/"+filename, myfile)
            content.memo = request.POST.get("memo", "")
            content.save()
            return redirect('/teacher/course/content/'+classroom_id+"/"+str(content.course_id))   
    return render(request,'teacher/course_edit.html',{'content': instance,  'content_id':content_id})		
	
def course_download(request, classroom_id, content_id):
    content = CourseContent.objects.get(id=content_id)
    filename = content.title
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))		
    download =  BASE_DIR + "/static/upload/" + content.filename
    wrapper = FileWrapper(file( download, "r" ))
    response = HttpResponse(wrapper, content_type = 'application/force-download')
    #response = HttpResponse(content_type='application/force-download')
    response['Content-Disposition'] = 'attachment; filename={0}'.format(filename.encode('utf8'))
    # It's usually a good idea to set the 'Content-Length' header too.
    # You can also set any other required headers: Cache-Control, etc.
    return response
    #return render(request,'student/download.html', {'download':download})
		
class CourseEditUpdateView(UpdateView):
    model = CourseWork
    fields = ['title']
    template_name = 'form.html'
    #success_url = '/teacher/course/domain/'
    def get_success_url(self):
        succ_url =  '/teacher/course/'+self.kwargs['classroom_id']
        return succ_url
	
def course_deadline(request, classroom_id, course_id):
    course = CourseWork.objects.get(id=course_id)
    classroom = Classroom.objects.get(id=classroom_id)
    if request.method == 'POST':
        form = CategroyForm(request.POST)
        if form.is_valid():
            course.domains = request.POST.getlist('domains')
            course.levels = request.POST.getlist('levels')	
            course.save()
            return redirect('/teacher/course/'+classroom_id+"/"+classroom_id)
    else:
        courseclass = CourseClass.objects.get(classroom_id=classroom_id, course_id=course_id)
        form = CourseDeadlineForm(instance=courseclass)
        courseclasses = CourseClass.objects.filter(course_id=course_id).order_by("-id")
    return render(request,'teacher/course_deadline_form.html',{'courseclasses':courseclasses, 'courseclass':courseclass, 'course':course, 'classroom':classroom})

	
# Ajax 設定期限、取消期限
def course_deadline_set(request):
    course_id = request.POST.get('courseid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        courseclass = CourseClass.objects.get(course_id=course_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        courseclass = Courseclass(course_id=course_id, classroom_id=classroom_id)
    if status == 'True':
        courseclass.deadline = True
    else :
        courseclass.deadline = False
    courseclass.save()
    return JsonResponse({'status':status}, safe=False)        

# Ajax 設定期限日期
def course_deadline_date(request):
    course_id = request.POST.get('courseid')
    classroom_id = request.POST.get('classroomid')		
    deadline_date = request.POST.get('deadlinedate')
    try:
        courseclass = CourseClass.objects.get(course_id=course_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        courseclass = CourseClass(course_id=course_id, classroom_id=classroom_id)
    #fclass.deadline_date = deadline_date.strftime('%d/%m/%Y')
    courseclass.deadline_date = datetime.strptime(deadline_date, '%Y %B %d - %H:%M')
    courseclass.save()
    return JsonResponse({'status':deadline_date}, safe=False)             

def course_group(request, classroom_id, course_id):
    enrolls = Enroll.objects.filter(classroom_id=classroom_id)
    enroll_dict = {}
    for enroll in enrolls:
        enroll_dict[enroll.id] = enroll
    groups = ClassroomGroup.objects.filter(classroom_id=classroom_id)
    group_ids = []
    for group in groups:
        group_ids.append(group.id)
    classroom = Classroom.objects.get(id=classroom_id)
    studentgroups = StudentGroup.objects.filter(group_id__in=group_ids)
    group_list = []
    for group in groups:        
        groupclass_list = []  
        groupclass_dict = {}
        students = filter(lambda student: student.group_id == group.id , studentgroups)
        for student in students:
            if student.enroll_id in enroll_dict:
                if student.group in groupclass_dict:
                    groupclass_dict[student.group].append(enroll_dict[student.enroll_id])
                else :
                    groupclass_dict[student.group] = [enroll_dict[student.enroll_id]]
        for key in groupclass_dict:
            groupclass_list.append([key, groupclass_dict[key]])
        group_list.append([group.id, groupclass_list])
    courseclass = CourseClass.objects.get(course_id=course_id, classroom_id=classroom_id)
    try:
        group = ClassroomGroup.objects.get(id=courseclass.group)
    except ObjectDoesNotExist:
        group = ClassroomGroup(title="不分組", id=0)
    return render(request,'teacher/course_group.html',{'course_id': course_id, 'coursegroup': group, 'groups':groups, 'classroom':classroom, 'group_list':group_list})

# Ajax 設定合作區組別
def course_group_set(request):
    course_id = request.POST.get('courseid')
    classroom_id = request.POST.get('classroomid')		
    group = request.POST.get('groupid')
    try:
        courseclass = CourseClass.objects.get(course_id=int(course_id), classroom_id=int(classroom_id))
    except ObjectDoesNotExist:
        pass
    courseclass.group = int(group)
    courseclass.save()    
    return JsonResponse({'status':course_id}, safe=False)  

# 列出所有練習主題
class CourseExerciseListView(ListView):
    model = CourseExercise
    context_object_name = 'exercises'
    template_name = "teacher/course_exercise.html"		
    paginate_by = 20
    def get_queryset(self):        
        queryset = CourseExercise.objects.filter(content_id=self.kwargs['content_id'])
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(CourseExerciseListView, self).get_context_data(**kwargs)
        content = CourseContent.objects.get(id=self.kwargs['content_id'])
        context['content'] = content
        context['classroom_id'] = self.kwargs['classroom_id']
        return context	

# 列出所有練習主題
class CourseExerciseAddListView(ListView):
    model = CourseExercise
    context_object_name = 'exercises'
    template_name = "teacher/course_exercise_add.html"		
    paginate_by = 20
    def get_queryset(self):
        #新增註記
        if self.kwargs['types'] ==  "0":        
            classes = SpeculationClass.objects.filter(classroom_id=self.kwargs['classroom_id'])
            speculation_ids = [speculation.forum_id for speculation in classes]
            works = SpeculationWork.objects.filter(id__in=speculation_ids).order_by("-id")            
        else:
            classes = ExamClass.objects.filter(classroom_id=self.kwargs['classroom_id'])
            exam_ids = [exam.exam_id for exam in classes]
            works = Exam.objects.filter(id__in=exam_ids).order_by("-id")
        exercises = CourseExercise.objects.filter(types=self.kwargs['types'], content_id=self.kwargs['content_id'])
        exercise_ids = [exercise.exercise_id for exercise in exercises]
        queryset = []
        for work in works:
            if work.id in exercise_ids:
                queryset.append([work, True])
            else:
                queryset.append([work, False])
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(CourseExerciseAddListView, self).get_context_data(**kwargs)
        content = CourseContent.objects.get(id=self.kwargs['content_id'])
        context['content'] = content
        context['types'] = self.kwargs['types']
        return context	

# 分組
def exercise_make(request):
    exercise_id = request.POST.get('exerciseid')
    content_id = request.POST.get('contentid')    
    types = request.POST.get('types') 
    action = request.POST.get('action')
    if content_id and exercise_id and action :      
        if action == 'set':
            try:
                exercise = CourseExercise.objects.get(types=types, content_id=content_id, exercise_id=exercise_id)  
            except ObjectDoesNotExist:          
                exercise = CourseExercise(types=types, content_id=content_id, exercise_id=exercise_id)  
                exercise.save()   
        else : 
            try:
                exercise = CourseExercise.objects.get(types=types, content_id=content_id, exercise_id=exercise_id)  
                exercise.delete()
            except ObjectDoesNotExist:          
                pass         
        return JsonResponse({'status':'ok'}, safe=False)
    else:
        return JsonResponse({'status':'fail'}, safe=False) 